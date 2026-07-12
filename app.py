import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import io
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from supabase import create_client, Client
import hashlib
import requests
import calendar
import bcrypt

# ============================================
# CONSTANTES DE GESTORES
# ============================================

GESTOR_MARCOS = "Sua Gestão - Chat Notas"
GESTOR_POLYANA = "Gestão Polyana Ventura - Chat Outros"
GESTORES_VALIDOS = [GESTOR_MARCOS, GESTOR_POLYANA]

# ============================================
# CONFIGURACOES INICIAIS
# ============================================

st.set_page_config(
    page_title="Sistema de Performance",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================
# SUPABASE CONFIG
# ============================================

def init_supabase():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except:
        url = os.environ.get("SUPABASE_URL")
        key = os.environ.get("SUPABASE_KEY")
        if url and key:
            return create_client(url, key)
        return None

# ============================================
# FUNCOES DE PADRONIZACAO
# ============================================

def padronizar_periodo(periodo):
    if not periodo:
        return periodo
    return periodo.strip().title()

def forcar_perfil_correto():
    if not st.session_state.get('logado', False):
        return
    
    usuario = st.session_state.get('usuario', '')
    
    if usuario == 'marcos':
        st.session_state.acesso_total = False
        st.session_state.perfil = "Gestor"
        st.session_state.gestor = GESTOR_MARCOS
    elif usuario == 'polyana':
        st.session_state.acesso_total = False
        st.session_state.perfil = "Gestor"
        st.session_state.gestor = GESTOR_POLYANA
    elif usuario == 'carine':
        st.session_state.acesso_total = True
        st.session_state.perfil = "Coordenador"
        st.session_state.gestor = GESTOR_MARCOS

def limpar_cache_completo():
    st.cache_data.clear()
    st.cache_resource.clear()
    
    keys_to_clear = ['resultados', 'processado', 'periodo', 'df_historico']
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    forcar_perfil_correto()
    st.success("✅ Cache limpo e perfil forçado!")
    st.rerun()

# ============================================
# FUNCOES DE RESET
# ============================================

def resetar_usuario_carine():
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não conectado")
        return False
    
    try:
        supabase.table('usuarios').delete().eq('usuario', 'carine').execute()
        supabase.table('usuarios').insert({
            'usuario': 'carine',
            'nome': 'Carine Melo',
            'senha': hash_senha('carine2026'),
            'gestor': GESTOR_MARCOS,
            'acesso_total': True
        }).execute()
        st.success("✅ Usuário Carine resetado com sucesso!")
        return True
    except Exception as e:
        st.error(f"❌ Erro ao resetar: {str(e)}")
        return False

def resetar_usuario_marcos():
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não conectado")
        return False
    
    try:
        supabase.table('usuarios').delete().eq('usuario', 'marcos').execute()
        supabase.table('usuarios').insert({
            'usuario': 'marcos',
            'nome': 'Marcos Miranda',
            'senha': hash_senha('marcos2026'),
            'gestor': GESTOR_MARCOS,
            'acesso_total': False
        }).execute()
        st.success("✅ Usuário Marcos resetado com sucesso!")
        return True
    except Exception as e:
        st.error(f"❌ Erro ao resetar: {str(e)}")
        return False

def resetar_usuario_polyana():
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não conectado")
        return False
    
    try:
        supabase.table('usuarios').delete().eq('usuario', 'polyana').execute()
        supabase.table('usuarios').insert({
            'usuario': 'polyana',
            'nome': 'Polyana Ventura',
            'senha': hash_senha('polyana2026'),
            'gestor': GESTOR_POLYANA,
            'acesso_total': False
        }).execute()
        st.success("✅ Usuário Polyana resetado com sucesso!")
        return True
    except Exception as e:
        st.error(f"❌ Erro ao resetar: {str(e)}")
        return False

def adicionar_dados_teste_polyana():
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não conectado")
        return False
    
    try:
        existing = supabase.table('historico_performance').select('*').eq('gestor', GESTOR_POLYANA).eq('mes_ano', 'Maio 2026').execute()
        if existing.data:
            st.info("ℹ️ Dados para Polyana já existem.")
            return True
        
        analistas_polyana = [
            'Christian Matozinho', 'Diego Machado', 'Igor Siqueira',
            'Ismael Chagas Bessa', 'João Pedro Santana', 'Karolyne Moreira',
            'Luan Pereira', 'Mario Junior', 'Maycon Oliveira',
            'Miguel Augusto', 'Polliana Santana'
        ]
        
        dados_teste = []
        for i, analista in enumerate(analistas_polyana):
            csat_base = 85 + (i * 0.7) % 12
            avaliacoes_base = 22 + (i * 0.5) % 8
            atendimentos_base = 130 + (i * 3) % 30
            meta_csat = 90 if analista in ['João Pedro Santana', 'Mario Junior'] else 86
            
            if csat_base >= meta_csat and avaliacoes_base >= 25:
                status = "🟢 Meta Superada"
            elif csat_base >= meta_csat or avaliacoes_base >= 25:
                status = "🟡 Atenção"
            else:
                status = "🔴 Crítico"
            
            dados_teste.append({
                'mes_ano': 'Maio 2026',
                'analista': analista,
                'gestor': GESTOR_POLYANA,
                'csat': round(csat_base, 1),
                'perc_avaliacoes': round(avaliacoes_base, 1),
                'perc_envio': round(100 - avaliacoes_base, 1),
                'total_atendimentos': int(atendimentos_base),
                'total_inativos': 3 + (i % 5),
                'validos': int(atendimentos_base) - (3 + (i % 5)),
                'avaliacoes': int(round(avaliacoes_base / 100 * atendimentos_base)),
                'positivos': int(round(csat_base / 100 * (avaliacoes_base / 100 * atendimentos_base))),
                'negativos': int(round((100 - csat_base) / 100 * (avaliacoes_base / 100 * atendimentos_base))),
                'meta_csat': meta_csat,
                'delta_csat': round(csat_base - meta_csat, 1),
                'meta_geral': 25,
                'status': status
            })
        
        for dado in dados_teste:
            supabase.table('historico_performance').insert(dado).execute()
        
        st.success(f"✅ Dados para {len(dados_teste)} analistas da Polyana adicionados!")
        return True
    except Exception as e:
        st.error(f"❌ Erro: {str(e)}")
        return False

def diagnosticar_sistema():
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não conectado")
        return
    
    st.subheader("🔍 Diagnóstico do Sistema")
    
    try:
        response = supabase.table('usuarios').select('*').execute()
        st.write("📋 Usuários no Supabase:")
        for user in response.data:
            st.write(f"- **{user['usuario']}** ({user['nome']}) - Acesso Total: {user.get('acesso_total', False)} - Gestor: {user['gestor']}")
    except Exception as e:
        st.warning(f"⚠️ Erro ao buscar usuários: {str(e)}")
    
    st.write("📋 Configuração dos Analistas:")
    analistas_config = carregar_analistas()
    for gestor, config in analistas_config.items():
        st.write(f"**{gestor}** - {len(config['membros'])} analistas")

# ============================================
# GERENCIAR USUARIOS
# ============================================

def gerenciar_usuarios_supabase():
    st.header("👥 Gerenciar Usuários no Supabase")
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não conectado")
        return
    
    try:
        response = supabase.table('usuarios').select('*').execute()
        usuarios = response.data
        
        if usuarios:
            st.subheader("📋 Usuários Cadastrados")
            dados_tabela = []
            for u in usuarios:
                dados_tabela.append({
                    'Usuário': u['usuario'],
                    'Nome': u['nome'],
                    'Gestor': u['gestor'],
                    'Perfil': 'Coordenador' if u.get('acesso_total') else 'Gestor',
                    'Ativo': '✅ Sim' if u.get('ativo', True) else '❌ Não'
                })
            st.dataframe(pd.DataFrame(dados_tabela), use_container_width=True, hide_index=True)
        
        st.markdown("---")
        
        st.subheader("➕ Criar Novo Usuário")
        col1, col2, col3 = st.columns(3)
        with col1:
            novo_usuario = st.text_input("Usuário (login)", key="novo_usuario_login")
            novo_nome = st.text_input("Nome completo", key="novo_usuario_nome")
        with col2:
            nova_senha = st.text_input("Senha", type="password", key="novo_usuario_senha")
            confirma_senha = st.text_input("Confirmar senha", type="password", key="novo_usuario_confirma")
        with col3:
            novo_gestor = st.selectbox("Gestor", GESTORES_VALIDOS, key="novo_usuario_gestor")
            novo_perfil = st.selectbox("Perfil", ["Gestor", "Coordenador"], key="novo_usuario_perfil")
            novo_ativo = st.checkbox("Ativo", value=True, key="novo_usuario_ativo")
        
        if st.button("✅ Cadastrar Usuário", use_container_width=True, key="btn_criar_usuario"):
            if not novo_usuario or not novo_nome or not nova_senha:
                st.error("❌ Preencha todos os campos!")
            elif novo_usuario in [u['usuario'] for u in usuarios]:
                st.error("❌ Usuário já existe!")
            elif nova_senha != confirma_senha:
                st.error("❌ Senhas não conferem!")
            elif len(nova_senha) < 6:
                st.error("❌ Senha deve ter no mínimo 6 caracteres!")
            else:
                senha_hash = hash_senha(nova_senha)
                acesso_total = True if novo_perfil == "Coordenador" else False
                try:
                    supabase.table('usuarios').insert({
                        'usuario': novo_usuario,
                        'nome': novo_nome,
                        'senha': senha_hash,
                        'gestor': novo_gestor,
                        'acesso_total': acesso_total,
                        'ativo': novo_ativo
                    }).execute()
                    st.success(f"✅ Usuário {novo_usuario} criado com sucesso!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Erro ao criar usuário: {str(e)}")
        
        st.markdown("---")
        
        st.subheader("✏️ Editar Usuário")
        if usuarios:
            usuario_selecionado = st.selectbox(
                "Selecione um usuário para editar",
                [u['usuario'] for u in usuarios],
                key="editar_usuario_select"
            )
            
            if usuario_selecionado:
                user_data = next((u for u in usuarios if u['usuario'] == usuario_selecionado), None)
                if user_data:
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        edit_nome = st.text_input("Nome", value=user_data['nome'], key="edit_usuario_nome")
                        edit_gestor = st.selectbox("Gestor", GESTORES_VALIDOS, 
                            index=GESTORES_VALIDOS.index(user_data['gestor']) if user_data['gestor'] in GESTORES_VALIDOS else 0,
                            key="edit_usuario_gestor")
                    with col2:
                        edit_senha = st.text_input("Nova senha (deixe em branco para manter)", type="password", key="edit_usuario_senha")
                        edit_confirma = st.text_input("Confirmar nova senha", type="password", key="edit_usuario_confirma")
                    with col3:
                        edit_perfil = st.selectbox("Perfil", ["Gestor", "Coordenador"],
                            index=1 if user_data.get('acesso_total') else 0,
                            key="edit_usuario_perfil")
                        edit_ativo = st.checkbox("Ativo", value=user_data.get('ativo', True), key="edit_usuario_ativo")
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        if st.button("💾 Salvar Alterações", use_container_width=True, key="btn_salvar_usuario"):
                            updates = {
                                'nome': edit_nome,
                                'gestor': edit_gestor,
                                'acesso_total': True if edit_perfil == "Coordenador" else False,
                                'ativo': edit_ativo
                            }
                            if edit_senha and edit_senha == edit_confirma and len(edit_senha) >= 6:
                                updates['senha'] = hash_senha(edit_senha)
                            elif edit_senha and edit_senha != edit_confirma:
                                st.error("❌ Senhas não conferem!")
                            else:
                                try:
                                    supabase.table('usuarios').update(updates).eq('usuario', usuario_selecionado).execute()
                                    st.success(f"✅ Usuário {usuario_selecionado} atualizado!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Erro ao atualizar: {str(e)}")
                    
                    with col2:
                        if st.button("🗑️ Excluir Usuário", use_container_width=True, key="btn_excluir_usuario"):
                            confirmar = st.checkbox("Confirmar exclusão permanente", key="confirmar_excluir_usuario")
                            if confirmar:
                                try:
                                    supabase.table('usuarios').delete().eq('usuario', usuario_selecionado).execute()
                                    st.success(f"✅ Usuário {usuario_selecionado} excluído!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Erro ao excluir: {str(e)}")
                            else:
                                st.warning("⚠️ Marque a caixa de confirmação para excluir.")
        
        st.markdown("---")
        if st.button("🔙 Voltar", key="voltar_usuarios"):
            st.session_state.gerenciar_usuarios = False
            st.rerun()
    except Exception as e:
        st.error(f"❌ Erro: {str(e)}")

# ============================================
# FUNCOES DE LEITURA DE ARQUIVOS
# ============================================

def normalizar_nome_coluna(nome):
    import re
    nome = nome.lower().strip()
    nome = re.sub(r'[^a-záéíóúãõç0-9\s]', '', nome)
    nome = re.sub(r'\s+', ' ', nome).strip()
    return nome

def identificar_coluna_flexivel(df, padroes):
    colunas = df.columns.tolist()
    for col in colunas:
        col_normalizado = normalizar_nome_coluna(col)
        for padrao in padroes:
            padrao_normalizado = normalizar_nome_coluna(padrao)
            if padrao_normalizado in col_normalizado or col_normalizado in padrao_normalizado:
                return col
    return None

def carregar_arquivo_satisfacao(arquivo):
    try:
        df = pd.read_excel(arquivo)
    except:
        try:
            df = pd.read_csv(arquivo, encoding='utf-8')
        except:
            try:
                df = pd.read_csv(arquivo, encoding='latin-1')
            except Exception as e:
                st.error(f"❌ Não foi possível ler o arquivo de satisfação: {str(e)}")
                return None
    
    colunas = df.columns.tolist()
    st.info(f"🔍 Colunas encontradas: {', '.join(colunas)}")
    
    col_id = identificar_coluna_flexivel(df, ['ID do ticket', 'Ticket ID', 'ID', 'ticket_id'])
    col_satisfacao = identificar_coluna_flexivel(df, ['Índice de satisfação do ticket', 'Satisfaction', 'satisfacao', 'satisfação'])
    col_atribuido = identificar_coluna_flexivel(df, ['Nome do atribuído', 'Assignee', 'Atribuído', 'Responsável'])
    
    if not col_id or not col_satisfacao or not col_atribuido:
        st.error("❌ Colunas necessárias não encontradas")
        return None
    
    df_renomeado = df.rename(columns={
        col_id: 'ID do ticket',
        col_satisfacao: 'Índice de satisfação do ticket',
        col_atribuido: 'Nome do atribuído'
    })
    
    colunas_necessarias = ['ID do ticket', 'Índice de satisfação do ticket', 'Nome do atribuído']
    df_final = df_renomeado[colunas_necessarias].copy()
    
    mapa_conversao = {
        'Boa': 'Good', 'boa': 'Good', 'Ruim': 'Bad', 'ruim': 'Bad',
        'Oferecida': 'Offered', 'oferecida': 'Offered',
        'Não oferecida': 'Unoffered', 'não oferecida': 'Unoffered',
        'Nao oferecida': 'Unoffered', 'nao oferecida': 'Unoffered'
    }
    df_final['Índice de satisfação do ticket'] = df_final['Índice de satisfação do ticket'].replace(mapa_conversao)
    
    st.success(f"✅ Arquivo carregado! {len(df_final)} registros")
    return df_final

def carregar_arquivo_inativos(arquivo):
    try:
        df = pd.read_excel(arquivo)
    except:
        try:
            df = pd.read_csv(arquivo, encoding='utf-8')
        except:
            try:
                df = pd.read_csv(arquivo, encoding='latin-1')
            except Exception as e:
                st.error(f"❌ Não foi possível ler o arquivo de inativos: {str(e)}")
                return None
    
    colunas = df.columns.tolist()
    st.info(f"🔍 Colunas encontradas: {', '.join(colunas)}")
    
    col_id = identificar_coluna_flexivel(df, ['ID do ticket', 'Ticket ID', 'ID', 'ticket_id'])
    col_atribuido = identificar_coluna_flexivel(df, ['Nome do atribuído', 'Assignee', 'Atribuído', 'Responsável'])
    
    if not col_id or not col_atribuido:
        st.error("❌ Colunas necessárias não encontradas")
        return None
    
    df_renomeado = df.rename(columns={
        col_id: 'ID do ticket',
        col_atribuido: 'Nome do atribuído'
    })
    
    colunas_necessarias = ['ID do ticket', 'Nome do atribuído']
    df_final = df_renomeado[colunas_necessarias].copy()
    
    st.success(f"✅ Arquivo carregado! {len(df_final)} registros")
    return df_final

# ============================================
# FUNCOES DE BANCO DE DADOS
# ============================================

def __tratar_erro_unique(e):
    erro_str = str(e).lower()
    if 'duplicate key' in erro_str or 'unique constraint' in erro_str or '23505' in erro_str:
        return True, "Já existe um registro para este analista neste período e gestor. Nenhum dado foi duplicado."
    return False, None

def salvar_historico(supabase, dados, mes_ano, gestor):
    if not supabase:
        return False, "Supabase não configurado"
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        dados_por_gestor = {}
        for analista, d in dados.items():
            gestor_real = d.get('gestor', gestor)
            if gestor_real not in dados_por_gestor:
                dados_por_gestor[gestor_real] = []
            dados_por_gestor[gestor_real].append((analista, d))
        
        for gestor_real, registros_do_gestor in dados_por_gestor.items():
            existing = supabase.table('historico_performance').select('*').eq('mes_ano', mes_ano).eq('gestor', gestor_real).execute()
            if existing.data:
                return False, f"Já existe relatório para {mes_ano} do gestor {gestor_real}"
            
            for analista, d in registros_do_gestor:
                existing_registro = supabase.table('historico_performance').select('*').eq('mes_ano', mes_ano).eq('analista', analista).eq('gestor', gestor_real).execute()
                if existing_registro.data:
                    return False, f"Já existe registro para {analista} em {mes_ano} do gestor {gestor_real}"
                
                try:
                    supabase.table('historico_performance').insert({
                        'mes_ano': mes_ano,
                        'analista': analista,
                        'gestor': gestor_real,
                        'csat': d['csat'],
                        'perc_avaliacoes': d['perc_avaliacoes'],
                        'perc_envio': d['perc_envio'],
                        'total_atendimentos': d['total_atendimentos'],
                        'total_inativos': d['total_inativos'],
                        'validos': d['validos'],
                        'avaliacoes': d['avaliacoes'],
                        'positivos': d['positivos'],
                        'negativos': d['negativos'],
                        'meta_csat': d['meta_csat'],
                        'delta_csat': d['delta_csat'],
                        'meta_geral': d['meta_geral'],
                        'status': d['status']
                    }).execute()
                except Exception as insert_error:
                    is_unique, msg = __tratar_erro_unique(insert_error)
                    if is_unique:
                        return False, f"Já existe registro para {analista} em {mes_ano} do gestor {gestor_real}. {msg}"
                    raise insert_error
        
        return True, "Salvo com sucesso"
    except Exception as e:
        return False, str(e)

def substituir_historico(supabase, dados, mes_ano, gestor):
    if not supabase:
        return False, "Supabase não configurado"
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        dados_por_gestor = {}
        for analista, d in dados.items():
            gestor_real = d.get('gestor', gestor)
            if gestor_real not in dados_por_gestor:
                dados_por_gestor[gestor_real] = []
            dados_por_gestor[gestor_real].append((analista, d))
        
        for gestor_real, registros_do_gestor in dados_por_gestor.items():
            try:
                delete_result = supabase.table('historico_performance').delete().eq('mes_ano', mes_ano).eq('gestor', gestor_real).execute()
                if hasattr(delete_result, 'error') and delete_result.error:
                    return False, f"Erro ao excluir registros antigos do gestor {gestor_real}: {delete_result.error}"
            except Exception as delete_error:
                return False, f"Erro ao excluir registros antigos do gestor {gestor_real}: {str(delete_error)}"
            
            for analista, d in registros_do_gestor:
                try:
                    supabase.table('historico_performance').insert({
                        'mes_ano': mes_ano,
                        'analista': analista,
                        'gestor': gestor_real,
                        'csat': d['csat'],
                        'perc_avaliacoes': d['perc_avaliacoes'],
                        'perc_envio': d['perc_envio'],
                        'total_atendimentos': d['total_atendimentos'],
                        'total_inativos': d['total_inativos'],
                        'validos': d['validos'],
                        'avaliacoes': d['avaliacoes'],
                        'positivos': d['positivos'],
                        'negativos': d['negativos'],
                        'meta_csat': d['meta_csat'],
                        'delta_csat': d['delta_csat'],
                        'meta_geral': d['meta_geral'],
                        'status': d['status']
                    }).execute()
                except Exception as insert_error:
                    is_unique, msg = __tratar_erro_unique(insert_error)
                    if is_unique:
                        return False, f"Já existe registro para {analista} em {mes_ano} do gestor {gestor_real} (possível conflito durante substituição). {msg}"
                    raise insert_error
        
        return True, "Substituído com sucesso"
    except Exception as e:
        return False, str(e)

def carregar_historico(supabase, mes_ano=None, gestor=None, analista=None):
    if not supabase:
        return None
    
    if mes_ano:
        mes_ano = padronizar_periodo(mes_ano)
    
    try:
        query = supabase.table('historico_performance').select('*')
        if mes_ano:
            query = query.eq('mes_ano', mes_ano)
        if gestor:
            query = query.eq('gestor', gestor)
        if analista:
            query = query.eq('analista', analista)
        response = query.execute()
        return pd.DataFrame(response.data)
    except Exception as e:
        st.error(f"Erro ao carregar histórico: {str(e)}")
        return None

def listar_periodos(supabase, gestor=None):
    if not supabase:
        return []
    
    try:
        query = supabase.table('historico_performance').select('mes_ano, gestor, data_criacao')
        if gestor:
            query = query.eq('gestor', gestor)
        response = query.execute()
        
        if gestor and len(response.data) == 0:
            all_response = supabase.table('historico_performance').select('gestor').execute()
            gestores_existentes = set([r['gestor'] for r in all_response.data])
            if gestores_existentes:
                st.info(f"🔍 Gestor '{gestor}' não encontrado. Gestores disponíveis: {', '.join(gestores_existentes)}")
        
        periodos = {}
        for item in response.data:
            mes_ano = item['mes_ano']
            if mes_ano not in periodos:
                periodos[mes_ano] = {
                    'mes_ano': mes_ano,
                    'gestor': item['gestor'],
                    'data_criacao': item.get('data_criacao', datetime.now().isoformat())
                }
        return list(periodos.values())
    except Exception as e:
        st.error(f"Erro ao listar períodos: {e}")
        return []

def excluir_periodo(supabase, mes_ano, gestor):
    if not supabase:
        return False, "Supabase não configurado"
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        supabase.table('historico_performance').delete().eq('mes_ano', mes_ano).eq('gestor', gestor).execute()
        supabase.table('podio_manual').delete().eq('mes_ano', mes_ano).eq('gestor', gestor).execute()
        supabase.table('podio_exclusoes').delete().eq('mes_ano', mes_ano).eq('gestor', gestor).execute()
        
        st.session_state.processado = False
        if 'resultados' in st.session_state:
            del st.session_state.resultados
        
        st.rerun()
        return True, "Período excluído com sucesso"
    except Exception as e:
        return False, str(e)

def verificar_periodo_existente(supabase, mes_ano, gestor):
    if not supabase:
        return False
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        response = supabase.table('historico_performance').select('*').eq('mes_ano', mes_ano).eq('gestor', gestor).execute()
        return len(response.data) > 0
    except:
        return False

def carregar_podio_manual(supabase, mes_ano, gestor):
    if not supabase:
        return None
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        response = supabase.table('podio_manual').select('*').eq('mes_ano', mes_ano).eq('gestor', gestor).order('posicao').execute()
        if response.data:
            return [(d['analista'], d['csat'], d['atendimentos'], 0) for d in response.data]
        return None
    except:
        return None

def salvar_podio_manual(supabase, mes_ano, gestor, podio):
    if not supabase:
        return False, "Supabase não configurado"
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        supabase.table('podio_manual').delete().eq('mes_ano', mes_ano).eq('gestor', gestor).execute()
        for i, (nome, csat, atendimentos, _) in enumerate(podio, 1):
            supabase.table('podio_manual').insert({
                'mes_ano': mes_ano,
                'gestor': gestor,
                'posicao': i,
                'analista': nome,
                'csat': csat,
                'atendimentos': atendimentos
            }).execute()
        return True, "Salvo com sucesso"
    except Exception as e:
        return False, str(e)

def carregar_exclusoes_podio(supabase, mes_ano, gestor):
    if not supabase:
        return []
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        response = supabase.table('podio_exclusoes').select('analista').eq('mes_ano', mes_ano).eq('gestor', gestor).execute()
        if response.data:
            return [item['analista'] for item in response.data]
        return []
    except Exception as e:
        return []

def salvar_exclusao_podio(supabase, mes_ano, gestor, analista):
    if not supabase:
        return False
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        supabase.table('podio_exclusoes').insert({
            'mes_ano': mes_ano,
            'gestor': gestor,
            'analista': analista
        }).execute()
        return True
    except Exception as e:
        if 'duplicate key' in str(e).lower() or 'unique constraint' in str(e).lower():
            return True
        return False

def remover_exclusao_podio(supabase, mes_ano, gestor, analista):
    if not supabase:
        return False
    
    mes_ano = padronizar_periodo(mes_ano)
    
    try:
        supabase.table('podio_exclusoes').delete().eq('mes_ano', mes_ano).eq('gestor', gestor).eq('analista', analista).execute()
        return True
    except Exception as e:
        return False

# ============================================
# FUNCOES DE BACKUP E RESTORE
# ============================================

def gerar_backup_excel(supabase):
    if not supabase:
        return None, "Supabase não configurado"
    
    try:
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            response = supabase.table('historico_performance').select('*').execute()
            df = pd.DataFrame(response.data)
            if not df.empty:
                df = df.drop(columns=['id'], errors='ignore')
            df.to_excel(writer, sheet_name='historico_performance', index=False)
            
            response = supabase.table('podio_manual').select('*').execute()
            df = pd.DataFrame(response.data)
            if not df.empty:
                df = df.drop(columns=['id'], errors='ignore')
            df.to_excel(writer, sheet_name='podio_manual', index=False)
            
            response = supabase.table('podio_exclusoes').select('*').execute()
            df = pd.DataFrame(response.data)
            if not df.empty:
                df = df.drop(columns=['id'], errors='ignore')
            df.to_excel(writer, sheet_name='podio_exclusoes', index=False)
        
        buffer.seek(0)
        return buffer, "Backup gerado com sucesso!"
    
    except Exception as e:
        return None, f"Erro ao gerar backup: {str(e)}"

def restaurar_backup_excel(supabase, arquivo_excel, modo):
    if not supabase:
        return False, "Supabase não configurado"
    
    try:
        excel_data = pd.read_excel(arquivo_excel, sheet_name=None)
        
        abas_esperadas = ['historico_performance', 'podio_manual', 'podio_exclusoes']
        for aba in abas_esperadas:
            if aba not in excel_data:
                return False, f"Arquivo de backup inválido: aba '{aba}' não encontrada"
        
        total_restaurados = 0
        total_pulados = 0
        
        chaves_unicas = {
            'historico_performance': ['mes_ano', 'analista', 'gestor'],
            'podio_manual': ['mes_ano', 'gestor', 'posicao'],
            'podio_exclusoes': ['mes_ano', 'gestor', 'analista']
        }
        
        for aba_nome, df in excel_data.items():
            if df.empty:
                continue
            
            df = df.drop(columns=['id'], errors='ignore')
            
            if modo == 'substituir':
                try:
                    supabase.table(aba_nome).delete().neq('id', 0).execute()
                except Exception as e:
                    response = supabase.table(aba_nome).select('id').execute()
                    for item in response.data:
                        supabase.table(aba_nome).delete().eq('id', item['id']).execute()
            
            registros = df.to_dict('records')
            chaves = chaves_unicas.get(aba_nome, [])
            
            for registro in registros:
                if modo == 'mesclar' and chaves:
                    query = supabase.table(aba_nome).select('*')
                    for chave in chaves:
                        if chave in registro:
                            query = query.eq(chave, registro[chave])
                    response = query.execute()
                    
                    if response.data:
                        total_pulados += 1
                        continue
                
                try:
                    supabase.table(aba_nome).insert(registro).execute()
                    total_restaurados += 1
                except Exception as e:
                    is_unique, _ = __tratar_erro_unique(e)
                    if is_unique:
                        total_pulados += 1
                    else:
                        raise e
        
        if modo == 'substituir':
            mensagem = f"Substituição concluída: {total_restaurados} registros restaurados, {total_pulados} falhas"
        else:
            mensagem = f"Mesclagem concluída: {total_restaurados} registros inseridos, {total_pulados} registros ignorados (já existentes)"
        
        return True, mensagem
    
    except Exception as e:
        return False, f"Erro ao restaurar backup: {str(e)}"

# ============================================
# GERENCIAMENTO DE USUARIOS
# ============================================

def hash_senha(senha):
    return hashlib.sha256(senha.encode()).hexdigest()

def carregar_usuarios():
    supabase = init_supabase()
    if supabase:
        try:
            response = supabase.table('usuarios').select('*').execute()
            if response.data:
                usuarios = {}
                for u in response.data:
                    usuarios[u['usuario']] = {
                        'senha': u['senha'],
                        'nome': u['nome'],
                        'gestor': u['gestor'],
                        'acesso_total': u.get('acesso_total', False),
                        'ativo': u.get('ativo', True)
                    }
                return usuarios
        except Exception as e:
            st.warning(f"⚠️ Erro ao carregar do Supabase: {str(e)}")
    
    try:
        if os.path.exists('usuarios.json'):
            with open('usuarios.json', 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    
    usuarios = {
        "marcos": {
            "senha": hash_senha("marcos2026"),
            "nome": "Marcos Miranda",
            "gestor": GESTOR_MARCOS,
            "acesso_total": False,
            "ativo": True
        },
        "polyana": {
            "senha": hash_senha("polyana2026"),
            "nome": "Polyana Ventura",
            "gestor": GESTOR_POLYANA,
            "acesso_total": False,
            "ativo": True
        },
        "carine": {
            "senha": hash_senha("carine2026"),
            "nome": "Carine Melo",
            "gestor": GESTOR_MARCOS,
            "acesso_total": True,
            "ativo": True
        }
    }
    
    try:
        with open('usuarios.json', 'w', encoding='utf-8') as f:
            json.dump(usuarios, f, ensure_ascii=False, indent=2)
    except:
        pass
    
    return usuarios

def salvar_usuario_supabase(supabase, usuario, nome, senha_hash, gestor, acesso_total=False):
    if not supabase:
        return False
    
    try:
        existing = supabase.table('usuarios').select('*').eq('usuario', usuario).execute()
        if existing.data:
            supabase.table('usuarios').update({
                'nome': nome,
                'senha': senha_hash,
                'gestor': gestor,
                'acesso_total': acesso_total
            }).eq('usuario', usuario).execute()
        else:
            supabase.table('usuarios').insert({
                'usuario': usuario,
                'nome': nome,
                'senha': senha_hash,
                'gestor': gestor,
                'acesso_total': acesso_total
            }).execute()
        return True
    except Exception as e:
        print(f"Erro ao salvar usuário: {e}")
        return False

# ============================================
# GERENCIAMENTO DE ANALISTAS
# ============================================

def carregar_analistas():
    try:
        if os.path.exists('analistas.json'):
            with open('analistas.json', 'r', encoding='utf-8') as f:
                config = json.load(f)
                if GESTOR_POLYANA in config and len(config[GESTOR_POLYANA]['membros']) >= 11:
                    return config
    except:
        pass
    
    config = {
        GESTOR_MARCOS: {
            "gestor": "Marcos Miranda",
            "meta_geral_avaliacoes": 25,
            "membros": {
                "Ana Claudia Corrêa": {"meta_csat": 90, "ativo": True},
                "João Vitor Almeida": {"meta_csat": 90, "ativo": True},
                "João Pedro Vianey": {"meta_csat": 90, "ativo": True},
                "Carlos Lemos": {"meta_csat": 90, "ativo": True},
                "Lorena Almeida": {"meta_csat": 86, "ativo": True},
                "Paulo Victor": {"meta_csat": 86, "ativo": True},
                "Rayane Nunes": {"meta_csat": 86, "ativo": True},
                "Thiago Reis": {"meta_csat": 90, "ativo": True},
                "Vanessa Silva": {"meta_csat": 86, "ativo": True}
            }
        },
        GESTOR_POLYANA: {
            "gestor": "Polyana Ventura",
            "meta_geral_avaliacoes": 25,
            "membros": {
                "Christian Matozinho": {"meta_csat": 86, "ativo": True},
                "Diego Machado": {"meta_csat": 86, "ativo": True},
                "Igor Siqueira": {"meta_csat": 86, "ativo": True},
                "Ismael Chagas Bessa": {"meta_csat": 86, "ativo": True},
                "João Pedro Santana": {"meta_csat": 90, "ativo": True},
                "Karolyne Moreira": {"meta_csat": 86, "ativo": True},
                "Luan Pereira": {"meta_csat": 86, "ativo": True},
                "Mario Junior": {"meta_csat": 90, "ativo": True},
                "Maycon Oliveira": {"meta_csat": 86, "ativo": True},
                "Miguel Augusto": {"meta_csat": 86, "ativo": True},
                "Polliana Santana": {"meta_csat": 86, "ativo": True}
            }
        }
    }
    
    try:
        with open('analistas.json', 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    except:
        pass
    
    return config

def salvar_analistas(analistas):
    try:
        with open('analistas.json', 'w', encoding='utf-8') as f:
            json.dump(analistas, f, ensure_ascii=False, indent=2)
        return True
    except:
        return False

# ============================================
# FUNCAO DE LOGIN
# ============================================

def fazer_login_corrigido():
    st.sidebar.markdown("---")
    st.sidebar.subheader("🔐 Login")
    
    supabase = init_supabase()
    if not supabase:
        st.sidebar.error("❌ Erro de conexão com o banco")
        return False
    
    usuario = st.sidebar.text_input("Usuário")
    senha = st.sidebar.text_input("Senha", type="password")
    
    if st.sidebar.button("Entrar", use_container_width=True):
        if not usuario or not senha:
            st.sidebar.error("❌ Preencha usuário e senha!")
            return False
        
        try:
            response = supabase.table('usuarios').select('*').eq('usuario', usuario).execute()
            
            if not response.data:
                st.sidebar.error("❌ Usuário não encontrado!")
                return False
            
            user = response.data[0]
            
            if not user.get('ativo', True):
                st.sidebar.error("❌ Usuário desativado!")
                return False
            
            try:
                import bcrypt
                senha_hash = user.get('senha', '')
                if bcrypt.checkpw(senha.encode('utf-8'), senha_hash.encode('utf-8')):
                    st.session_state.logado = True
                    st.session_state.usuario = user['usuario']
                    st.session_state.nome_usuario = user['nome']
                    st.session_state.gestor = user.get('gestor', user.get('time_nome', ''))
                    st.session_state.acesso_total = user.get('acesso_total', False)
                    st.session_state.perfil = "Coordenador" if user.get('acesso_total', False) else "Gestor"
                    st.session_state.cargo = user.get('cargo', '')
                    st.session_state.time_nome = user.get('time_nome', '')
                    st.session_state.supervisor_nome = user.get('supervisor_nome', '')
                    
                    try:
                        supabase.table('usuarios').update({
                            'ultimo_login': datetime.now().isoformat()
                        }).eq('id', user['id']).execute()
                    except:
                        pass
                    
                    forcar_perfil_correto()
                    
                    st.sidebar.success(f"✅ Bem-vindo, {user['nome']}!")
                    st.rerun()
                    return True
                else:
                    st.sidebar.error("❌ Senha incorreta!")
                    return False
                    
            except ImportError:
                st.sidebar.error("❌ Erro ao verificar senha!")
                return False
                
        except Exception as e:
            st.sidebar.error(f"❌ Erro: {str(e)}")
            return False
    
    if st.session_state.get('logado', False):
        st.sidebar.markdown("---")
        st.sidebar.subheader("👤 Perfil")
        st.sidebar.write(f"**Usuário:** {st.session_state.get('nome_usuario')}")
        st.sidebar.write(f"**Perfil:** {st.session_state.get('perfil')}")
        if st.session_state.get('acesso_total', False):
            st.sidebar.write("**Acesso:** 🔑 Total")
        else:
            st.sidebar.write(f"**Time:** {st.session_state.get('gestor')}")
        
        st.sidebar.success("✅ Logado")
        
        if st.sidebar.button("Sair", use_container_width=True):
            keys_to_clear = ['logado', 'usuario', 'nome_usuario', 'gestor', 'acesso_total', 
                           'perfil', 'cargo', 'time_nome', 'supervisor_nome', 'resultados', 
                           'processado', 'periodo', 'df_historico']
            for key in keys_to_clear:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()
        return True
    
    return False

# ============================================
# FUNCOES AUXILIARES
# ============================================

def get_genero_neutro(nome):
    nomes_femininos = [
        'Ana', 'Maria', 'Paula', 'Lorena', 'Vanessa', 'Rayane', 
        'Karolyne', 'Polliana', 'Polyana', 'Ariane', 'Julia',
        'Claudia', 'Almeida', 'Carine'
    ]
    for nome_feminino in nomes_femininos:
        if nome.startswith(nome_feminino):
            return "colaboradora"
    if nome.endswith('a') and not nome.endswith('as'):
        return "colaboradora"
    return "colaborador"

def processar_dados(df_satisfacao, df_inativos, analistas_config):
    resultados = {}
    
    analistas_ativos = []
    for gestor, config in analistas_config.items():
        for analista, dados in config['membros'].items():
            if dados['ativo']:
                analistas_ativos.append(analista)
    
    for analista in analistas_ativos:
        tickets_analista = df_satisfacao[df_satisfacao['Nome do atribuído'] == analista]
        total_atendimentos = len(tickets_analista)
        if total_atendimentos == 0:
            continue
        
        inativos_analista = df_inativos[df_inativos['Nome do atribuído'] == analista]
        total_inativos = len(inativos_analista)
        validos = total_atendimentos - total_inativos
        
        avaliacoes = len(tickets_analista[tickets_analista['Índice de satisfação do ticket'].isin(['Good', 'Bad'])])
        positivos = len(tickets_analista[tickets_analista['Índice de satisfação do ticket'] == 'Good'])
        negativos = len(tickets_analista[tickets_analista['Índice de satisfação do ticket'] == 'Bad'])
        
        perc_avaliacoes = (avaliacoes / validos * 100) if validos > 0 else 0
        csat = (positivos / avaliacoes * 100) if avaliacoes > 0 else 0
        nao_avaliaram = validos - avaliacoes
        perc_envio = (nao_avaliaram / validos * 100) if validos > 0 else 0
        
        gestor = "Outros"
        meta_csat = 86
        meta_geral = 25
        
        for gestor_nome, config in analistas_config.items():
            if analista in config['membros']:
                gestor = gestor_nome
                meta_csat = config['membros'][analista]['meta_csat']
                meta_geral = config.get('meta_geral_avaliacoes', 25)
                break
        
        delta_csat = csat - meta_csat
        
        if csat >= meta_csat and perc_avaliacoes >= meta_geral:
            status = "🟢 Meta Superada"
        elif csat >= meta_csat or perc_avaliacoes >= meta_geral:
            status = "🟡 Atenção"
        else:
            status = "🔴 Crítico"
        
        resultados[analista] = {
            'total_atendimentos': total_atendimentos,
            'total_inativos': total_inativos,
            'validos': validos,
            'avaliacoes': avaliacoes,
            'positivos': positivos,
            'negativos': negativos,
            'perc_avaliacoes': round(perc_avaliacoes, 2),
            'perc_envio': round(perc_envio, 2),
            'csat': round(csat, 2),
            'meta_csat': meta_csat,
            'delta_csat': round(delta_csat, 2),
            'meta_geral': meta_geral,
            'status': status,
            'gestor': gestor
        }
    
    return resultados

def calcular_media_operacao(resultados):
    if not resultados:
        return 0
    total_atendimentos = sum([d['total_atendimentos'] for d in resultados.values()])
    total_analistas = len(resultados)
    return round(total_atendimentos / total_analistas) if total_analistas > 0 else 0

def calcular_podio(resultados, media_atendimentos=None, limiar_csat=90):
    if not resultados:
        return []
    
    if media_atendimentos is None:
        media_atendimentos = calcular_media_operacao(resultados)
    
    dados_validos = {k: v for k, v in resultados.items() 
                     if v['csat'] >= limiar_csat 
                     and v['total_atendimentos'] >= media_atendimentos 
                     and v['perc_avaliacoes'] >= 25}
    
    if not dados_validos:
        return []
    
    sorted_analistas = sorted(dados_validos.items(), key=lambda x: x[1]['csat'], reverse=True)
    top_3 = sorted_analistas[:3]
    return [(nome, dados['csat'], dados['total_atendimentos'], dados['perc_avaliacoes']) for nome, dados in top_3]

def ordenar_meses(meses):
    ordem_meses = {
        'January': 1, 'February': 2, 'March': 3, 'April': 4,
        'May': 5, 'June': 6, 'July': 7, 'August': 8,
        'September': 9, 'October': 10, 'November': 11, 'December': 12,
        'Janeiro': 1, 'Fevereiro': 2, 'Março': 3, 'Abril': 4,
        'Maio': 5, 'Junho': 6, 'Julho': 7, 'Agosto': 8,
        'Setembro': 9, 'Outubro': 10, 'Novembro': 11, 'Dezembro': 12
    }
    
    def extrair_ano_mes(mes_str):
        if not mes_str:
            return (9999, 13)
        partes = mes_str.split()
        if len(partes) >= 2:
            nome_mes = partes[0]
            try:
                ano = int(partes[1])
            except ValueError:
                ano = 9999
            numero_mes = ordem_meses.get(nome_mes, 13)
            return (ano, numero_mes)
        return (9999, 13)
    
    return sorted(meses, key=extrair_ano_mes)

# ============================================
# FUNCOES DE CALCULO - TOP PERFORMERS E OPORTUNIDADES
# ============================================

def calcular_top_performers(resultados, media_atendimentos):
    """
    Calcula os Top Performers considerando:
    - CSAT ≥ 90%
    - Atendimentos ≥ média da equipe
    - Avaliações ≥ 25%
    """
    if not resultados:
        return []
    
    top_performers = []
    for analista, dados in resultados.items():
        # Critérios para Top Performer
        csat_ok = dados['csat'] >= 90
        atendimentos_ok = dados['total_atendimentos'] >= media_atendimentos
        avaliacoes_ok = dados['perc_avaliacoes'] >= 25
        
        if csat_ok and atendimentos_ok and avaliacoes_ok:
            top_performers.append({
                'analista': analista,
                'csat': dados['csat'],
                'atendimentos': dados['total_atendimentos'],
                'avaliacoes': dados['perc_avaliacoes'],
                'status': dados['status']
            })
    
    # Ordenar por CSAT (decrescente)
    top_performers.sort(key=lambda x: x['csat'], reverse=True)
    return top_performers[:5]

def calcular_oportunidades(resultados, media_atendimentos):
    """
    Calcula as Oportunidades considerando:
    - CSAT abaixo da meta individual
    - OU avaliações abaixo de 25%
    - OU atendimentos abaixo da média
    """
    if not resultados:
        return []
    
    oportunidades = []
    for analista, dados in resultados.items():
        # Critérios para Oportunidade
        csat_abaixo = dados['csat'] < dados['meta_csat']
        avaliacoes_abaixo = dados['perc_avaliacoes'] < 25
        atendimentos_abaixo = dados['total_atendimentos'] < media_atendimentos
        
        # Se atender a PELO MENOS UM critério, é oportunidade
        if csat_abaixo or avaliacoes_abaixo or atendimentos_abaixo:
            # Identificar motivo principal
            motivos = []
            if csat_abaixo:
                motivos.append(f"CSAT abaixo da meta ({dados['csat']:.1f}% < {dados['meta_csat']}%)")
            if avaliacoes_abaixo:
                motivos.append(f"Avaliações abaixo de 25% ({dados['perc_avaliacoes']:.1f}%)")
            if atendimentos_abaixo:
                motivos.append(f"Atendimentos abaixo da média ({dados['total_atendimentos']} < {media_atendimentos:.0f})")
            
            oportunidades.append({
                'analista': analista,
                'csat': dados['csat'],
                'meta_csat': dados['meta_csat'],
                'delta_csat': dados['delta_csat'],
                'atendimentos': dados['total_atendimentos'],
                'avaliacoes': dados['perc_avaliacoes'],
                'motivos': motivos,
                'status': dados['status']
            })
    
    # Ordenar por CSAT (crescente - piores primeiro)
    oportunidades.sort(key=lambda x: x['csat'])
    return oportunidades[:5]

# ============================================
# FUNCOES DE DASHBOARD VISUAL - OTIMIZADAS
# ============================================

def criar_card_compacto(titulo, valor, meta, subtitulo, cor):
    """Card compacto para uso em grids"""
    st.markdown(f"""
    <div style="background: var(--background-color, #ffffff); 
                padding: 12px; 
                border-radius: 8px; 
                border-left: 4px solid {cor};
                margin-bottom: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.05);">
        <div style="display: flex; justify-content: space-between; align-items: center;">
            <span style="font-size: 12px; color: var(--text-secondary, #666);">{titulo}</span>
            <span style="font-size: 20px; font-weight: bold; color: var(--text-color, #333);">{valor}</span>
        </div>
        <div style="display: flex; justify-content: space-between; margin-top: 4px;">
            <span style="font-size: 10px; color: var(--text-secondary, #999);">Meta: {meta}</span>
            <span style="font-size: 10px; color: var(--text-secondary, #999);">{subtitulo}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

def criar_cards_indicadores_compactos(dados, meta_csat=90, meta_avaliacoes=25, meta_envio=90):
    """Cards superiores em layout compacto"""
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        criar_card_compacto(
            "📊 Analistas",
            dados.get('total_registros', 0),
            "-",
            "Total",
            "#2ecc71"
        )
    
    with col2:
        cor_csat = "#2ecc71" if dados.get('csat_medio', 0) >= meta_csat else "#e74c3c"
        criar_card_compacto(
            "⭐ CSAT",
            f"{dados.get('csat_medio', 0):.1f}%",
            f"{meta_csat}%",
            "Médio",
            cor_csat
        )
    
    with col3:
        cor_avaliacoes = "#2ecc71" if dados.get('perc_avaliacoes_medio', 0) >= meta_avaliacoes else "#e74c3c"
        criar_card_compacto(
            "📝 Avaliações",
            f"{dados.get('perc_avaliacoes_medio', 0):.1f}%",
            f"{meta_avaliacoes}%",
            "Médio",
            cor_avaliacoes
        )
    
    with col4:
        cor_envio = "#2ecc71" if dados.get('perc_envio_medio', 0) >= meta_envio else "#e74c3c"
        criar_card_compacto(
            "📤 Envio",
            f"{dados.get('perc_envio_medio', 0):.1f}%",
            f"{meta_envio}%",
            "Médio",
            cor_envio
        )

def criar_painel_inteligente_compacto(csat_medio, perc_avaliacoes_medio, perc_envio_medio, 
                                    metas_superadas, total_analistas):
    """Painel inteligente em layout compacto - 2 colunas"""
    
    col1, col2, col3, col4 = st.columns(4)
    
    saude_csat = "🟢" if csat_medio >= 90 else "🟡" if csat_medio >= 85 else "🔴"
    saude_avaliacoes = "🟢" if perc_avaliacoes_medio >= 25 else "🟡" if perc_avaliacoes_medio >= 20 else "🔴"
    saude_envio = "🟢" if perc_envio_medio >= 80 else "🟡" if perc_envio_medio >= 70 else "🔴"
    
    with col1:
        criar_card_compacto(
            "🏆 Metas Superadas",
            f"{metas_superadas}/{total_analistas}",
            "-",
            "Analistas",
            "#2ecc71"
        )
    
    with col2:
        criar_card_compacto(
            "⭐ CSAT",
            f"{saude_csat} {csat_medio:.1f}%",
            "Meta: 90%",
            "Status",
            "#2ecc71" if csat_medio >= 90 else "#f39c12" if csat_medio >= 85 else "#e74c3c"
        )
    
    with col3:
        criar_card_compacto(
            "📝 Avaliações",
            f"{saude_avaliacoes} {perc_avaliacoes_medio:.1f}%",
            "Meta: 25%",
            "Status",
            "#2ecc71" if perc_avaliacoes_medio >= 25 else "#f39c12" if perc_avaliacoes_medio >= 20 else "#e74c3c"
        )
    
    with col4:
        criar_card_compacto(
            "📤 Envio",
            f"{saude_envio} {perc_envio_medio:.1f}%",
            "Meta: 80%",
            "Status",
            "#2ecc71" if perc_envio_medio >= 80 else "#f39c12" if perc_envio_medio >= 70 else "#e74c3c"
        )
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🔔 Alertas")
        alertas = []
        
        if csat_medio < 85:
            alertas.append("🚨 CSAT crítico (< 85%)")
        elif csat_medio < 90:
            alertas.append("⚡ CSAT em atenção")
        else:
            alertas.append("✅ CSAT saudável")
        
        if perc_avaliacoes_medio < 20:
            alertas.append("🚨 Avaliações críticas (< 20%)")
        elif perc_avaliacoes_medio < 25:
            alertas.append("⚡ Avaliações abaixo da meta")
        else:
            alertas.append("✅ Avaliações no alvo")
        
        if perc_envio_medio < 70:
            alertas.append("🚨 Envio crítico (< 70%)")
        elif perc_envio_medio < 80:
            alertas.append("⚡ Envio abaixo da meta")
        else:
            alertas.append("✅ Envio saudável")
        
        for alerta in alertas:
            st.markdown(f"- {alerta}")
    
    with col2:
        st.markdown("#### 💡 Recomendações")
        recomendacoes = []
        
        if csat_medio >= 90 and perc_avaliacoes_medio >= 25:
            recomendacoes.append("🏆 Excelente! Mantenha o padrão")
        elif csat_medio >= 90:
            recomendacoes.append("⭐ Foco em aumentar avaliações")
        elif perc_avaliacoes_medio >= 25:
            recomendacoes.append("📊 Foco em elevar CSAT")
        else:
            recomendacoes.append("🎯 Prioridade: CSAT e Avaliações")
        
        for recomendacao in recomendacoes:
            st.markdown(f"- {recomendacao}")

def criar_grafico_evolucao_unificado(df_historico, titulo="Evolução dos Indicadores"):
    """Gráfico unificado com todas as métricas"""
    if df_historico is None or df_historico.empty:
        return None
    
    meses_ordenados = ordenar_meses(df_historico['mes_ano'].unique().tolist())
    df_ordenado = df_historico.copy()
    df_ordenado['ordem'] = df_ordenado['mes_ano'].apply(lambda x: meses_ordenados.index(x) if x in meses_ordenados else 999)
    df_ordenado = df_ordenado.sort_values('ordem')
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=df_ordenado['mes_ano'],
        y=df_ordenado['csat'],
        mode='lines+markers',
        name='CSAT',
        line=dict(color='#2ecc71', width=3),
        marker=dict(size=8, color='#2ecc71'),
        hovertemplate='<b>%{x}</b><br>CSAT: %{y:.2f}%<extra></extra>'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_ordenado['mes_ano'],
        y=df_ordenado['perc_avaliacoes'],
        mode='lines+markers',
        name='% Avaliações',
        line=dict(color='#3498db', width=3),
        marker=dict(size=8, color='#3498db'),
        hovertemplate='<b>%{x}</b><br>Avaliações: %{y:.2f}%<extra></extra>'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_ordenado['mes_ano'],
        y=df_ordenado['perc_envio'],
        mode='lines+markers',
        name='% Envio',
        line=dict(color='#f39c12', width=3),
        marker=dict(size=8, color='#f39c12'),
        hovertemplate='<b>%{x}</b><br>Envio: %{y:.2f}%<extra></extra>'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_ordenado['mes_ano'],
        y=[90] * len(df_ordenado),
        mode='lines',
        name='Meta CSAT',
        line=dict(color='#e74c3c', width=2, dash='dash'),
        hovertemplate='Meta CSAT: 90%<extra></extra>'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_ordenado['mes_ano'],
        y=[25] * len(df_ordenado),
        mode='lines',
        name='Meta Avaliações',
        line=dict(color='#9b59b6', width=2, dash='dash'),
        hovertemplate='Meta Avaliações: 25%<extra></extra>'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_ordenado['mes_ano'],
        y=[80] * len(df_ordenado),
        mode='lines',
        name='Meta Envio',
        line=dict(color='#1abc9c', width=2, dash='dash'),
        hovertemplate='Meta Envio: 80%<extra></extra>'
    ))
    
    fig.update_layout(
        title=titulo,
        height=400,
        xaxis_title='Período',
        yaxis_title='Percentual (%)',
        yaxis_range=[0, 100],
        hovermode='x unified',
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=-0.25,
            xanchor='center',
            x=0.5
        ),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='var(--text-color, #e0e0e0)'),
        xaxis=dict(
            gridcolor='rgba(128,128,128,0.15)',
            tickfont=dict(color='var(--text-color, #e0e0e0)')
        ),
        yaxis=dict(
            gridcolor='rgba(128,128,128,0.15)',
            tickfont=dict(color='var(--text-color, #e0e0e0)')
        )
    )
    
    return fig

def gerar_grafico_mensal(analista, df_historico, meta_csat, meta_avaliacoes):
    """Gráfico de evolução individual"""
    if df_historico is None or df_historico.empty:
        return None
    
    df_analista = df_historico[df_historico['analista'] == analista]
    if df_analista.empty or len(df_analista['mes_ano'].unique()) < 2:
        return None
    
    df_analista = df_analista.sort_values('mes_ano')
    meses_ordenados = ordenar_meses(df_analista['mes_ano'].unique().tolist())
    df_analista['mes_ano_ordem'] = df_analista['mes_ano'].apply(lambda x: meses_ordenados.index(x) if x in meses_ordenados else 999)
    df_analista = df_analista.sort_values('mes_ano_ordem')
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=df_analista['mes_ano'],
        y=df_analista['csat'],
        mode='lines+markers',
        name='CSAT',
        line=dict(color='#2ecc71', width=3),
        marker=dict(size=10, color='#2ecc71')
    ))
    
    fig.add_trace(go.Bar(
        x=df_analista['mes_ano'],
        y=df_analista['perc_avaliacoes'],
        name='% Avaliações',
        marker_color='#3498db',
        opacity=0.6,
        yaxis='y2'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_analista['mes_ano'],
        y=[meta_csat]*len(df_analista),
        name=f'Meta CSAT: {meta_csat}%',
        line=dict(color='#e74c3c', width=2, dash='dash'),
        mode='lines'
    ))
    
    fig.add_trace(go.Scatter(
        x=df_analista['mes_ano'],
        y=[meta_avaliacoes]*len(df_analista),
        name=f'Meta Avaliações: {meta_avaliacoes}%',
        line=dict(color='#f39c12', width=2, dash='dot'),
        mode='lines',
        yaxis='y2'
    ))
    
    fig.update_layout(
        title=f"📈 Evolução Mensal - {analista}",
        height=380,
        xaxis_title='Período',
        yaxis=dict(
            title='CSAT (%)',
            range=[0, 100],
            side='left',
            gridcolor='rgba(128,128,128,0.15)'
        ),
        yaxis2=dict(
            title='% Avaliações',
            range=[0, 100],
            overlaying='y',
            side='right',
            gridcolor='rgba(128,128,128,0.05)'
        ),
        hovermode='x unified',
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=-0.25,
            xanchor='center',
            x=0.5
        ),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='var(--text-color, #e0e0e0)')
    )
    
    return fig

def mostrar_podio(podio):
    """Exibe o pódio em layout compacto"""
    if not podio:
        st.info("🏆 Nenhum analista no pódio este mês")
        return
    
    col1, col2, col3 = st.columns(3)
    medalhas = ["🥇", "🥈", "🥉"]
    cores = ['#FFD700', '#C0C0C0', '#CD7F32']
    
    for i, (col, (nome, csat, atendimentos, perc_avaliacoes)) in enumerate(zip([col1, col2, col3], podio), 1):
        with col:
            st.markdown(f"""
            <div style="text-align: center; padding: 15px; border: 2px solid var(--border-color, #444); 
                        border-radius: 10px; background: var(--background-color, rgba(255,255,255,0.05));
                        color: var(--text-color, #e0e0e0);">
                <div style="font-size: 36px;">{medalhas[i-1]}</div>
                <h4 style="margin: 5px 0; color: var(--text-color, #fff);">{i}º Lugar</h4>
                <h3 style="margin: 5px 0; color: var(--text-color, #fff);">{nome}</h3>
                <div style="font-size: 24px; font-weight: bold; color: {cores[i-1]};">{csat:.1f}%</div>
                <div style="font-size: 12px; color: var(--text-secondary, #ccc);">💬 {atendimentos} atendimentos</div>
                <div style="font-size: 12px; color: #28a745;">{perc_avaliacoes:.1f}% avaliações</div>
            </div>
            """, unsafe_allow_html=True)

# ============================================
# FUNCOES DE FEEDBACK
# ============================================

def gerar_analise_tecnica(analista, dados, media_operacao, podio):
    genero = get_genero_neutro(analista)
    posicao = "não está no pódio"
    for i, (nome, csat, atendimentos, perc_avaliacoes) in enumerate(podio, 1):
        if nome == analista:
            posicao = f"está em {i}º lugar no pódio com CSAT de {csat:.2f}%, {atendimentos} atendimentos e {perc_avaliacoes:.2f}% de avaliações"
            break
    
    if dados['delta_csat'] > 0:
        analise_csat = f"superou a meta em {dados['delta_csat']:.2f} pontos percentuais"
    elif dados['delta_csat'] < 0:
        analise_csat = f"ficou abaixo da meta em {abs(dados['delta_csat']):.2f} pontos percentuais"
    else:
        analise_csat = "atingiu exatamente a meta"
    
    diff_avaliacoes = dados['perc_avaliacoes'] - dados['meta_geral']
    if diff_avaliacoes >= 0:
        analise_avaliacoes = f"superou a meta de {dados['meta_geral']:.0f}% em {diff_avaliacoes:.2f} pontos percentuais"
    else:
        analise_avaliacoes = f"ficou abaixo da meta de {dados['meta_geral']:.0f}% em {abs(diff_avaliacoes):.2f} pontos percentuais"
    
    if dados['total_atendimentos'] > media_operacao:
        produtividade = f"{(dados['total_atendimentos'] / media_operacao - 1) * 100:.2f}% superior à média da operação"
    elif dados['total_atendimentos'] < media_operacao:
        produtividade = f"{(1 - dados['total_atendimentos'] / media_operacao) * 100:.2f}% inferior à média da operação"
    else:
        produtividade = "igual à média da operação"
    
    if dados['csat'] >= 94:
        nivel_csat = "extremamente alto"
    elif dados['csat'] >= 90:
        nivel_csat = "alto"
    elif dados['csat'] >= 85:
        nivel_csat = "bom"
    else:
        nivel_csat = "precisa de atenção"
    
    return f"""
### 1. Qualidade e Satisfação do Cliente (CSAT)

O(A) {genero} registrou um índice de **Satisfação (CSAT) de {dados['csat']:.2f}%**.

- **Comparativo com a Meta:** O resultado {analise_csat} (meta: ≥ {dados['meta_csat']:.0f}%).
- **Análise Detalhada:** Do volume total de feedbacks recebidos ({dados['avaliacoes']}), **{dados['positivos']} foram positivos**, resultando em um índice de aprovação {nivel_csat}. Houve apenas {dados['negativos']} registros negativos.

### 2. Engajamento e Coleta de Feedback

O(A) {genero} alcançou uma **Taxa de Avaliações de {dados['perc_avaliacoes']:.2f}%**.

- **Comparativo com a Meta:** A meta esperada é de no mínimo {dados['meta_geral']:.0f}% de conversão de atendimentos em avaliações. O resultado {analise_avaliacoes}.
- **Cálculo:** A taxa foi calculada sobre o volume total de {dados['avaliacoes']} avaliações divididas pelos **{dados['validos']} atendimentos válidos**.

### 3. Produtividade e Volumetria

O volume total de atendimentos realizados pelo(a) {genero} foi de **{dados['total_atendimentos']} chamados**.

- **Comparativo com a Operação:** A média de atendimentos por agente foi de {media_operacao}. {analista} absorveu uma demanda operacional **{produtividade}**.
- **Destaque:** {posicao} do mês em CSAT.
"""

def gerar_feedback_manual(analista, dados, media_operacao, posicao_podio=None, modo_feedback="SARI - Desenvolvimento"):
    """Gera feedback manual baseado no modo selecionado"""
    genero = get_genero_neutro(analista)
    texto_podio = f"🏆 {posicao_podio}º lugar no pódio do mês!" if posicao_podio else ""
    
    if modo_feedback == "MIMO - Operacional":
        return _gerar_feedback_mimo(analista, dados, media_operacao, posicao_podio, genero, texto_podio)
    else:
        return _gerar_feedback_sari(analista, dados, media_operacao, posicao_podio, genero, texto_podio)

def _gerar_feedback_mimo(analista, dados, media_operacao, posicao_podio, genero, texto_podio):
    feedback = f"""
### 1. MANTER
"""
    if dados['status'] == "🟢 Meta Superada":
        feedback += f"Parabéns! Você manteve um excelente desempenho com CSAT de {dados['csat']:.2f}% e taxa de avaliações de {dados['perc_avaliacoes']:.2f}%.\n\n"
    elif dados['status'] == "🟡 Atenção":
        feedback += f"Manteve consistência com CSAT de {dados['csat']:.2f}% e {dados['total_atendimentos']} atendimentos realizados.\n\n"
    else:
        feedback += f"Manteve o engajamento com {dados['avaliacoes']} avaliações coletadas e {dados['total_atendimentos']} atendimentos.\n\n"
    
    feedback += "### 2. MELHORAR\n"
    if dados['csat'] < dados['meta_csat']:
        feedback += f"Elevar o CSAT de {dados['csat']:.2f}% para atingir a meta de {dados['meta_csat']:.0f}%\n\n"
    elif dados['perc_avaliacoes'] < dados['meta_geral']:
        feedback += f"Aumentar a taxa de avaliações de {dados['perc_avaliacoes']:.2f}% para alcançar {dados['meta_geral']:.0f}%\n\n"
    else:
        feedback += "Buscar excelência contínua e compartilhar boas práticas com a equipe\n\n"
    
    feedback += "### 3. MOTIVO\n"
    if dados['csat'] < dados['meta_csat']:
        feedback += f"Clientes insatisfeitos ({(100 - dados['csat']):.2f}%) podem impactar negativamente a imagem da marca.\n\n"
    elif dados['perc_avaliacoes'] < dados['meta_geral']:
        feedback += "A coleta de avaliações é essencial para entender a experiência do cliente e identificar oportunidades de melhoria.\n\n"
    else:
        feedback += "Manter o alto padrão de atendimento fortalece a confiança dos clientes e a reputação da equipe.\n\n"
    
    feedback += "### 4. ORIENTAÇÃO PRÁTICA\n"
    if dados['csat'] < dados['meta_csat']:
        feedback += "1. Revisar os atendimentos com avaliações negativas\n"
        feedback += "2. Identificar padrões de insatisfação nos atendimentos\n"
        feedback += "3. Aplicar melhorias no processo de atendimento"
    elif dados['perc_avaliacoes'] < dados['meta_geral']:
        feedback += "1. Oferecer a pesquisa ao final de cada atendimento\n"
        feedback += "2. Explicar ao cliente a importância do feedback\n"
        feedback += "3. Monitorar diariamente a taxa de coleta"
    else:
        feedback += "1. Compartilhar boas práticas com a equipe\n"
        feedback += "2. Participar de treinamentos avançados\n"
        feedback += "3. Manter o alto padrão de atendimento"
    
    feedback += f"\n\n---\n**Status:** {dados['status']}\n{texto_podio}\n**Data:** {datetime.now().strftime('%d/%m/%Y')}"
    
    return feedback

def _gerar_feedback_sari(analista, dados, media_operacao, posicao_podio, genero, texto_podio):
    feedback = f"""
### 1. SITUAÇÃO

No período analisado, o(a) {genero} {analista} apresentou CSAT de {dados['csat']:.2f}% (meta: {dados['meta_csat']:.0f}%), "
taxa de avaliações de {dados['perc_avaliacoes']:.2f}% (meta: {dados['meta_geral']:.0f}%), "
e {dados['total_atendimentos']} atendimentos realizados. "
O percentual de envio foi de {dados['perc_envio']:.2f}%. {texto_podio}

### 2. ANÁLISE

"""
    if dados['delta_csat'] > 0:
        feedback += f"**CSAT:** Superou a meta em {dados['delta_csat']:.2f} pontos percentuais.\n"
    elif dados['delta_csat'] < 0:
        feedback += f"**CSAT:** Ficou abaixo da meta em {abs(dados['delta_csat']):.2f} pontos percentuais.\n"
    else:
        feedback += "**CSAT:** Atingiu exatamente a meta.\n"
    
    diff_avaliacoes = dados['perc_avaliacoes'] - dados['meta_geral']
    if diff_avaliacoes >= 0:
        feedback += f"**Avaliações:** Superou a meta de {dados['meta_geral']:.0f}% em {diff_avaliacoes:.2f} pontos percentuais.\n\n"
    else:
        feedback += f"**Avaliações:** Ficou abaixo da meta de {dados['meta_geral']:.0f}% em {abs(diff_avaliacoes):.2f} pontos percentuais.\n\n"
    
    feedback += "### 3. RECOMENDAÇÃO\n"
    if dados['csat'] < dados['meta_csat']:
        feedback += "1. Realizar análise aprofundada dos atendimentos negativos\n"
        feedback += "2. Identificar padrões de insatisfação recorrentes\n"
        feedback += "3. Participar de treinamentos específicos para áreas de melhoria\n\n"
    elif dados['perc_avaliacoes'] < dados['meta_geral']:
        feedback += "1. Revisar a abordagem de oferecimento da pesquisa\n"
        feedback += "2. Treinar técnicas de engajamento para coleta de feedback\n"
        feedback += "3. Criar rotina de verificação da coleta de avaliações\n\n"
    else:
        feedback += "1. Compartilhar boas práticas com a equipe\n"
        feedback += "2. Buscar certificações e treinamentos avançados\n"
        feedback += "3. Atuar como mentor para novos colaboradores\n\n"
    
    feedback += "### 4. IMPACTO E EXPECTATIVA\n"
    feedback += "**Impacto:** "
    if dados['csat'] >= 90:
        feedback += "Alta qualidade no atendimento prestado, contribuindo positivamente para a satisfação do cliente.\n\n"
    else:
        feedback += "Oportunidades de melhoria identificadas no atendimento prestado.\n\n"
    
    feedback += "**Expectativa:** "
    if dados['csat'] >= dados['meta_csat']:
        feedback += "Manter o CSAT acima da meta, "
    else:
        feedback += f"Elevar o CSAT para atingir a meta de {dados['meta_csat']:.0f}%, "
    
    if dados['perc_avaliacoes'] >= dados['meta_geral']:
        feedback += "continuar com a boa taxa de avaliações, "
    else:
        feedback += f"aumentar a taxa de avaliações para {dados['meta_geral']:.0f}%, "
    
    feedback += "e continuar evoluindo seu desempenho global."
    
    feedback += f"\n\n---\n**Status:** {dados['status']}\n{texto_podio}\n**Data:** {datetime.now().strftime('%d/%m/%Y')}"
    
    return feedback

def gerar_feedback_ia(analista, dados, media_operacao, posicao_podio=None, feedback_editado=None, modo_feedback="SARI - Desenvolvimento"):
    """Gera feedback usando IA"""
    genero = get_genero_neutro(analista)
    texto_podio = f"🏆 {posicao_podio}º lugar no pódio do mês!" if posicao_podio else "Não está no pódio"
    
    if "MIMO" in modo_feedback:
        prompt = _construir_prompt_mimo(analista, dados, media_operacao, genero, texto_podio, feedback_editado)
    else:
        prompt = _construir_prompt_sari(analista, dados, media_operacao, genero, texto_podio, feedback_editado)
    
    try:
        github_token = st.secrets.get("GITHUB_TOKEN", os.environ.get("GITHUB_TOKEN", ""))
        if github_token:
            headers = {"Authorization": f"Bearer {github_token}", "Content-Type": "application/json"}
            url = "https://models.inference.ai.azure.com/chat/completions"
            payload = {
                "model": "gpt-4o-mini",
                "messages": [{"role": "system", "content": "Você é especialista em gestão de performance."}, {"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 1500
            }
            response = requests.post(url, headers=headers, json=payload)
            if response.status_code == 200:
                result = response.json()
                return result['choices'][0]['message']['content'].strip()
        
        return gerar_feedback_manual(analista, dados, media_operacao, posicao_podio, modo_feedback)
    except Exception as e:
        return gerar_feedback_manual(analista, dados, media_operacao, posicao_podio, modo_feedback)

def _construir_prompt_mimo(analista, dados, media_operacao, genero, texto_podio, feedback_editado):
    base_prompt = f"""
Você é um especialista em gestão de performance e desenvolvimento de equipes de atendimento ao cliente.

## DADOS DO COLABORADOR:
- Nome: {analista}
- Gênero: {genero}
- CSAT: {dados['csat']:.2f}% (Meta: ≥ {dados['meta_csat']:.0f}%)
- Delta CSAT: {dados['delta_csat']:+.2f} pontos
- % Avaliações: {dados['perc_avaliacoes']:.2f}% (Meta: ≥ {dados['meta_geral']:.0f}%)
- % Envio: {dados['perc_envio']:.2f}%
- Atendimentos: {dados['total_atendimentos']}
- Média da operação: {media_operacao}
- Avaliações: {dados['avaliacoes']} ({dados['positivos']} positivas, {dados['negativos']} negativas)
- Status: {dados['status']}
- Posição no pódio: {texto_podio}

## ESTRUTURA OBRIGATÓRIA DO FEEDBACK (MIMO - OPERACIONAL):
### 1. MANTER (o que está bom e deve ser mantido)
### 2. MELHORAR (o que precisa evoluir)
### 3. MOTIVO (por que é importante melhorar)
### 4. ORIENTAÇÃO PRÁTICA (ações concretas)

## CARACTERÍSTICAS DO FEEDBACK MIMO:
- Feedback curto e objetivo
- Foco operacional
- Linguagem simples e direta
- Ações práticas e executáveis
- Tom profissional mas acessível
"""
    
    if feedback_editado:
        return base_prompt + f"""
## FEEDBACK ATUAL (EDITADO PELO GESTOR):
{feedback_editado}

## INSTRUÇÃO:
Analise o feedback acima e melhore/ajuste seguindo a estrutura MIMO obrigatória.
Gere o feedback revisado e aprimorado:
"""
    else:
        return base_prompt + """
## INSTRUÇÃO:
Gere um feedback completo de performance seguindo a estrutura MIMO (Manter, Melhorar, Motivo, Orientação).
O feedback deve ser curto, objetivo e focado em ações práticas.

Gere o feedback:
"""

def _construir_prompt_sari(analista, dados, media_operacao, genero, texto_podio, feedback_editado):
    base_prompt = f"""
Você é um especialista em gestão de performance e desenvolvimento de equipes de atendimento ao cliente.

## DADOS DO COLABORADOR:
- Nome: {analista}
- Gênero: {genero}
- CSAT: {dados['csat']:.2f}% (Meta: ≥ {dados['meta_csat']:.0f}%)
- Delta CSAT: {dados['delta_csat']:+.2f} pontos
- % Avaliações: {dados['perc_avaliacoes']:.2f}% (Meta: ≥ {dados['meta_geral']:.0f}%)
- % Envio: {dados['perc_envio']:.2f}%
- Atendimentos: {dados['total_atendimentos']}
- Média da operação: {media_operacao}
- Avaliações: {dados['avaliacoes']} ({dados['positivos']} positivas, {dados['negativos']} negativas)
- Status: {dados['status']}
- Posição no pódio: {texto_podio}

## ESTRUTURA OBRIGATÓRIA DO FEEDBACK (SARI - DESENVOLVIMENTO):
### 1. SITUAÇÃO (contexto atual do colaborador)
### 2. ANÁLISE (causas e fatores do desempenho)
### 3. RECOMENDAÇÃO (ações específicas para evolução)
### 4. IMPACTO E EXPECTATIVA (consequências e resultados esperados)

## CARACTERÍSTICAS DO FEEDBACK SARI:
- Feedback aprofundado e estruturado
- Foco em desenvolvimento de pessoas
- Conexão entre comportamento e resultado
- Linguagem profissional e construtiva
- Visão de longo prazo
"""
    
    if feedback_editado:
        return base_prompt + f"""
## FEEDBACK ATUAL (EDITADO PELO GESTOR):
{feedback_editado}

## INSTRUÇÃO:
Analise o feedback acima e melhore/ajuste seguindo a estrutura SARI obrigatória.
Gere o feedback revisado e aprimorado:
"""
    else:
        return base_prompt + """
## INSTRUÇÃO:
Gere um feedback completo de performance seguindo a estrutura SARI (Situação, Análise, Recomendação, Impacto e Expectativa).
O feedback deve ser aprofundado, construtivo e focado no desenvolvimento do colaborador.

Gere o feedback:
"""

def gerar_relatorio_word(analista, dados, analise_tecnica, feedback, media_operacao, podio, periodo):
    doc = Document()
    titulo = doc.add_heading(f'Relatório de Performance - {analista}', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Período: {periodo}')
    doc.add_paragraph('')
    doc.add_heading('Esperado:', level=1)
    doc.add_paragraph(f'≥ {dados["meta_geral"]:.0f}% de avaliações', style='List Bullet')
    doc.add_paragraph(f'≥ {dados["meta_csat"]:.0f}% de Satisfação', style='List Bullet')
    doc.add_paragraph('')
    doc.add_heading('Atingido:', level=1)
    
    posicao_texto = "Não está no pódio"
    for i, (nome, csat, atendimentos, perc_avaliacoes) in enumerate(podio, 1):
        if nome == analista:
            posicao_texto = f"{i}º Lugar - CSAT: {csat:.2f}% | {atendimentos} atendimentos | {perc_avaliacoes:.2f}% avaliações"
            break
    
    doc.add_paragraph(f'CSAT: {dados["csat"]:.2f}%', style='List Bullet')
    doc.add_paragraph(f'Avaliações: {dados["perc_avaliacoes"]:.2f}% ({dados["positivos"]} positivos + {dados["negativos"]} negativos = {dados["avaliacoes"]})', style='List Bullet')
    doc.add_paragraph(f'% Envio: {dados["perc_envio"]:.2f}%', style='List Bullet')
    doc.add_paragraph(f'Atendidos: {dados["total_atendimentos"]} - {dados["total_inativos"]} = {dados["validos"]}', style='List Bullet')
    doc.add_paragraph(f'Média por agente: {media_operacao}', style='List Bullet')
    doc.add_paragraph(f'Posição no Pódio: {posicao_texto}', style='List Bullet')
    doc.add_paragraph('')
    doc.add_heading('Análise Técnica de Desempenho', level=1)
    
    for linha in analise_tecnica.split('\n'):
        if linha.strip():
            if linha.strip().startswith('### 1.'):
                doc.add_heading('Qualidade e Satisfação do Cliente (CSAT)', level=2)
            elif linha.strip().startswith('### 2.'):
                doc.add_heading('Engajamento e Coleta de Feedback', level=2)
            elif linha.strip().startswith('### 3.'):
                doc.add_heading('Produtividade e Volumetria', level=2)
            else:
                doc.add_paragraph(linha.strip())
    
    if feedback:
        doc.add_paragraph('')
        doc.add_heading('Feedback de Performance', level=1)
        doc.add_paragraph(f'Status Geral do Período: {dados["status"]}')
        doc.add_paragraph('')
        
        for linha in feedback.split('\n'):
            if linha.strip() and not linha.strip().startswith('#'):
                if linha.strip().startswith('###'):
                    doc.add_heading(linha.strip().replace('###', '').strip(), level=2)
                else:
                    doc.add_paragraph(linha.strip())
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================
# DASHBOARD GESTOR - VERSÃO OTIMIZADA
# ============================================

def dashboard_gestor_otimizado(periodo, gestor_nome, supabase):
    periodo = padronizar_periodo(periodo)
    
    df_historico = carregar_historico(supabase, mes_ano=periodo, gestor=gestor_nome)
    
    if df_historico is None or df_historico.empty:
        st.warning(f"⚠️ Nenhum dado encontrado para {gestor_nome} no período {periodo}")
        return None
    
    resultados = {}
    for _, row in df_historico.iterrows():
        resultados[row['analista']] = {
            'total_atendimentos': row['total_atendimentos'],
            'total_inativos': row['total_inativos'],
            'validos': row['validos'],
            'avaliacoes': row['avaliacoes'],
            'positivos': row['positivos'],
            'negativos': row['negativos'],
            'perc_avaliacoes': row['perc_avaliacoes'],
            'perc_envio': row['perc_envio'],
            'csat': row['csat'],
            'meta_csat': row['meta_csat'],
            'delta_csat': row['delta_csat'],
            'meta_geral': row['meta_geral'],
            'status': row['status'],
            'gestor': row['gestor']
        }
    
    total_analistas = len(resultados)
    total_atendimentos = sum([d['total_atendimentos'] for d in resultados.values()])
    media_atendimentos = total_atendimentos / total_analistas if total_analistas > 0 else 0
    csat_medio = sum([d['csat'] for d in resultados.values()]) / total_analistas if total_analistas > 0 else 0
    perc_avaliacoes_medio = sum([d['perc_avaliacoes'] for d in resultados.values()]) / total_analistas if total_analistas > 0 else 0
    perc_envio_medio = sum([d['perc_envio'] for d in resultados.values()]) / total_analistas if total_analistas > 0 else 0
    metas_superadas = len([d for d in resultados.values() if d['status'] == '🟢 Meta Superada'])
    
    dados_cards = {
        'total_registros': total_analistas,
        'csat_medio': csat_medio,
        'perc_avaliacoes_medio': perc_avaliacoes_medio,
        'perc_envio_medio': perc_envio_medio
    }
    
    st.info(f"📅 {periodo} | 👤 {gestor_nome} | 🏆 Metas: {metas_superadas}/{total_analistas}")
    st.markdown("---")
    
    criar_cards_indicadores_compactos(dados_cards)
    st.markdown("---")
    
    criar_painel_inteligente_compacto(
        csat_medio, perc_avaliacoes_medio, perc_envio_medio,
        metas_superadas, total_analistas
    )
    st.markdown("---")
    
    tab_geral, tab_ranking, tab_analise, tab_individual = st.tabs([
        "📊 Visão Geral",
        "🏆 Rankings",
        "📋 Análise Detalhada",
        "👤 Individual"
    ])
    
    # ===== TAB 1: VISÃO GERAL =====
    with tab_geral:
        st.subheader("📈 Evolução dos Indicadores")
        df_historico_completo = carregar_historico(supabase, gestor=gestor_nome)
        
        if df_historico_completo is not None and not df_historico_completo.empty and len(df_historico_completo['mes_ano'].unique()) >= 2:
            df_mensal = df_historico_completo.groupby('mes_ano').agg({
                'csat': 'mean',
                'perc_avaliacoes': 'mean',
                'perc_envio': 'mean'
            }).reset_index()
            
            fig_evolucao = criar_grafico_evolucao_unificado(df_mensal)
            if fig_evolucao:
                st.plotly_chart(fig_evolucao, use_container_width=True)
        else:
            st.info("📊 Dados insuficientes para evolução. São necessários pelo menos 2 meses.")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🏆 Top Performers")
            top_performers = calcular_top_performers(resultados, media_atendimentos)
            
            if top_performers:
                for i, performer in enumerate(top_performers[:3], 1):
                    medalha = ["🥇", "🥈", "🥉"][i-1]
                    cores = ['#FFD700', '#C0C0C0', '#CD7F32']
                    
                    criterios = []
                    if performer['csat'] >= 90:
                        criterios.append(f"✅ CSAT {performer['csat']:.1f}% ≥ 90%")
                    if performer['atendimentos'] >= media_atendimentos:
                        criterios.append(f"✅ {performer['atendimentos']} ≥ {media_atendimentos:.0f} (média)")
                    if performer['avaliacoes'] >= 25:
                        criterios.append(f"✅ {performer['avaliacoes']:.1f}% ≥ 25%")
                    
                    st.markdown(f"""
                    <div style="background: var(--background-color, #f0f8ff); 
                                padding: 12px; border-radius: 8px; margin-bottom: 8px;
                                border-left: 4px solid {cores[i-1]};
                                color: var(--text-color, #333);">
                        <div style="display: flex; justify-content: space-between; align-items: center;">
                            <b style="color: {cores[i-1]};">{medalha} {performer['analista']}</b>
                            <span style="font-size: 18px; font-weight: bold; color: {cores[i-1]};">
                                {performer['csat']:.1f}%
                            </span>
                        </div>
                        <div style="font-size: 11px; color: var(--text-secondary, #666); margin-top: 4px;">
                            {' | '.join(criterios)}
                        </div>
                        <div style="font-size: 11px; color: #28a745; margin-top: 2px;">
                            {performer['status']}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.info("Nenhum analista atingiu todos os critérios para Top Performer.")
        
        with col2:
            st.subheader("📊 Oportunidades de Melhoria")
            oportunidades = calcular_oportunidades(resultados, media_atendimentos)
            
            if oportunidades:
                for i, oport in enumerate(oportunidades[:3], 1):
                    cor = "#e74c3c" if oport['csat'] < 85 else "#f39c12"
                    
                    # Status real (não "Meta Superada")
                    if oport['csat'] < oport['meta_csat']:
                        status_real = "🔴 CSAT abaixo da meta"
                    elif oport['avaliacoes'] < 25:
                        status_real = "🟡 Avaliações abaixo de 25%"
                    elif oport['atendimentos'] < media_atendimentos:
                        status_real = "🟡 Atendimentos abaixo da média"
                    else:
                        status_real = "⚡ Atenção necessária"
                    
                    st.markdown(f"""
                    <div style="background: var(--background-color, #fff5f5); 
                                padding: 12px; border-radius: 8px; margin-bottom: 8px;
                                border-left: 4px solid {cor};
                                color: var(--text-color, #333);">
                        <div style="display: flex; justify-content: space-between; align-items: center;">
                            <b>{i}º {oport['analista']}</b>
                            <span style="font-size: 16px; font-weight: bold; color: {cor};">
                                {oport['csat']:.1f}%
                            </span>
                        </div>
                        <div style="font-size: 11px; color: var(--text-secondary, #666); margin-top: 4px;">
                            {' | '.join(oport['motivos'][:2])}
                        </div>
                        <div style="font-size: 11px; color: {cor}; margin-top: 2px;">
                            {status_real}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.info("🎯 Nenhuma oportunidade identificada! Equipe excelente!")
    
    # ===== TAB 2: RANKINGS =====
    with tab_ranking:
        st.subheader("🏆 Pódio do Mês")
        
        analistas_excluidos = carregar_exclusoes_podio(supabase, periodo, gestor_nome) if supabase else []
        resultados_para_podio = {k: v for k, v in resultados.items() if k not in analistas_excluidos}
        podio = calcular_podio(resultados_para_podio, media_atendimentos)
        
        podio_manual = carregar_podio_manual(supabase, periodo, gestor_nome) if supabase else None
        podio_efetivo = podio_manual if podio_manual else podio
        mostrar_podio(podio_efetivo)
        
        st.markdown("---")
        st.subheader("📊 Ranking Completo")
        dados_ranking = []
        for i, (analista, dados) in enumerate(sorted(resultados.items(), key=lambda x: x[1]['csat'], reverse=True), 1):
            is_top = (dados['csat'] >= 90 and 
                     dados['total_atendimentos'] >= media_atendimentos and 
                     dados['perc_avaliacoes'] >= 25)
            
            is_opp = (dados['csat'] < dados['meta_csat'] or 
                     dados['perc_avaliacoes'] < 25 or 
                     dados['total_atendimentos'] < media_atendimentos)
            
            dados_ranking.append({
                'Posição': i,
                'Analista': analista,
                'CSAT': f"{dados['csat']:.1f}%",
                'Meta CSAT': f"{dados['meta_csat']}%",
                'Delta CSAT': f"{dados['delta_csat']:+.1f}%",
                'Atendimentos': dados['total_atendimentos'],
                'Média': f"{media_atendimentos:.0f}",
                'Avaliações': f"{dados['perc_avaliacoes']:.1f}%",
                'Meta Aval.': f"{dados['meta_geral']}%",
                'Categoria': "🏆 Top" if is_top else "⚠️ Oportunidade" if is_opp else "✅ Bom"
            })
        
        st.dataframe(pd.DataFrame(dados_ranking), use_container_width=True, hide_index=True)
    
    # ===== TAB 3: ANÁLISE DETALHADA =====
    with tab_analise:
        st.subheader("📋 Análise de Desempenho")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            filtro_status = st.selectbox("Status", ["Todos", "🟢 Meta Superada", "🟡 Atenção", "🔴 Crítico"])
        with col2:
            filtro_analista = st.selectbox("Analista", ["Todos"] + list(resultados.keys()))
        with col3:
            filtro_ordenacao = st.selectbox("Ordenar por", ["CSAT (decrescente)", "CSAT (crescente)", "Avaliações", "Atendimentos"])
        
        dados_filtrados = resultados
        if filtro_status != "Todos":
            dados_filtrados = {k: v for k, v in dados_filtrados.items() if v['status'] == filtro_status}
        if filtro_analista != "Todos":
            dados_filtrados = {k: v for k, v in dados_filtrados.items() if k == filtro_analista}
        
        if filtro_ordenacao == "CSAT (decrescente)":
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['csat'], reverse=True)
        elif filtro_ordenacao == "CSAT (crescente)":
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['csat'])
        elif filtro_ordenacao == "Avaliações":
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['perc_avaliacoes'], reverse=True)
        else:
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['total_atendimentos'], reverse=True)
        
        dados_tabela = []
        for analista, dados in dados_ordenados:
            delta_avaliacoes = dados['perc_avaliacoes'] - dados['meta_geral']
            dados_tabela.append({
                'Analista': analista,
                'CSAT': f"{dados['csat']:.1f}%",
                'Meta CSAT': f"{dados['meta_csat']}%",
                'Delta CSAT': f"{dados['delta_csat']:+.1f}%",
                '% Avaliações': f"{dados['perc_avaliacoes']:.1f}%",
                'Meta Avaliações': f"{dados['meta_geral']}%",
                'Delta Avaliações': f"{delta_avaliacoes:+.1f}%",
                '% Envio': f"{dados['perc_envio']:.1f}%",
                '💬 Atendimentos': dados['total_atendimentos'],
                'Status': dados['status']
            })
        
        if dados_tabela:
            st.dataframe(pd.DataFrame(dados_tabela), use_container_width=True, hide_index=True)
            
            st.markdown("---")
            st.subheader("📊 Resumo Estatístico")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Média Atendimentos", f"{media_atendimentos:.0f}")
            with col2:
                st.metric("Maior Atendimentos", f"{max([d['total_atendimentos'] for d in resultados.values()])}")
            with col3:
                st.metric("Menor Atendimentos", f"{min([d['total_atendimentos'] for d in resultados.values()])}")
            with col4:
                st.metric("Total Atendimentos", f"{total_atendimentos}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Média CSAT", f"{csat_medio:.1f}%")
            with col2:
                st.metric("Média Avaliações", f"{perc_avaliacoes_medio:.1f}%")
            with col3:
                st.metric("Média Envio", f"{perc_envio_medio:.1f}%")
        else:
            st.info("Nenhum dado encontrado com os filtros selecionados.")
    
    # ===== TAB 4: INDIVIDUAL =====
    with tab_individual:
        st.subheader("👤 Relatório Individual")
        
        analista_selecionado = st.selectbox(
            "Selecione o Analista",
            list(resultados.keys()),
            key="analista_individual_select"
        )
        
        if analista_selecionado:
            dados = resultados[analista_selecionado]
            
            posicao_podio = None
            for i, (nome, _, _, _) in enumerate(podio_efetivo, 1):
                if nome == analista_selecionado:
                    posicao_podio = i
                    break
            
            # Cards do analista
            st.markdown(f"### 📊 {analista_selecionado}")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                cor_csat = "#2ecc71" if dados['csat'] >= dados['meta_csat'] else "#e74c3c"
                criar_card_compacto(
                    "⭐ CSAT",
                    f"{dados['csat']:.1f}%",
                    f"Meta: {dados['meta_csat']}%",
                    f"{dados['delta_csat']:+.1f}%",
                    cor_csat
                )
            
            with col2:
                diff_avaliacoes = dados['perc_avaliacoes'] - dados['meta_geral']
                cor_avaliacoes = "#2ecc71" if diff_avaliacoes >= 0 else "#e74c3c"
                criar_card_compacto(
                    "📝 Avaliações",
                    f"{dados['perc_avaliacoes']:.1f}%",
                    f"Meta: {dados['meta_geral']}%",
                    f"{diff_avaliacoes:+.1f}%",
                    cor_avaliacoes
                )
            
            with col3:
                diff_envio = dados['perc_envio'] - 80
                cor_envio = "#2ecc71" if diff_envio >= 0 else "#f39c12"
                criar_card_compacto(
                    "📤 Envio",
                    f"{dados['perc_envio']:.1f}%",
                    "Meta: 80%",
                    f"{diff_envio:+.1f}%",
                    cor_envio
                )
            
            with col4:
                diff_atend = dados['total_atendimentos'] - media_atendimentos
                cor_atend = "#2ecc71" if diff_atend >= 0 else "#e74c3c"
                criar_card_compacto(
                    "💬 Atendimentos",
                    f"{dados['total_atendimentos']}",
                    f"Média: {media_atendimentos:.0f}",
                    f"{diff_atend:+.0f}",
                    cor_atend
                )
            
            # Status e posição
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**Status:** {dados['status']}")
            with col2:
                if posicao_podio:
                    medalhas = ["🥇", "🥈", "🥉"]
                    st.markdown(f"**Posição:** {medalhas[posicao_podio-1]} {posicao_podio}º lugar no pódio")
                else:
                    st.markdown("**Posição:** Não está no pódio")
            
            st.markdown("---")
            
            # Radar e Evolução
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📊 Radar de Performance")
                categorias = ['CSAT', '% Avaliações', '% Envio']
                valores = [dados['csat'], dados['perc_avaliacoes'], dados['perc_envio']]
                metas = [dados['meta_csat'], dados['meta_geral'], 80]
                
                fig_radar = go.Figure()
                fig_radar.add_trace(go.Scatterpolar(
                    r=valores,
                    theta=categorias,
                    fill='toself',
                    name=analista_selecionado,
                    line_color='#2ecc71',
                    fillcolor='rgba(46, 204, 113, 0.3)'
                ))
                fig_radar.add_trace(go.Scatterpolar(
                    r=metas,
                    theta=categorias,
                    fill='toself',
                    name='Meta',
                    line_color='#e74c3c',
                    fillcolor='rgba(231, 76, 60, 0.1)',
                    line_dash='dash'
                ))
                fig_radar.update_layout(
                    height=380,
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 100],
                            tickfont=dict(color='var(--text-color, #e0e0e0)')
                        ),
                        angularaxis=dict(
                            tickfont=dict(color='var(--text-color, #e0e0e0)')
                        )
                    ),
                    showlegend=True,
                    legend=dict(
                        orientation='h',
                        yanchor='bottom',
                        y=-0.15,
                        xanchor='center',
                        x=0.5
                    ),
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='var(--text-color, #e0e0e0)')
                )
                st.plotly_chart(fig_radar, use_container_width=True)
            
            with col2:
                st.subheader("📈 Evolução Mensal")
                
                if supabase:
                    if st.session_state.get('acesso_total', False):
                        df_historico_analista = carregar_historico(supabase, analista=analista_selecionado)
                    else:
                        df_historico_analista = carregar_historico(supabase, analista=analista_selecionado, gestor=gestor_nome)
                    
                    if df_historico_analista is not None and not df_historico_analista.empty and len(df_historico_analista['mes_ano'].unique()) >= 2:
                        fig_individual = gerar_grafico_mensal(
                            analista_selecionado,
                            df_historico_analista,
                            dados['meta_csat'],
                            dados['meta_geral']
                        )
                        if fig_individual:
                            st.plotly_chart(fig_individual, use_container_width=True)
                        else:
                            st.info("📊 Dados insuficientes para gráfico individual.")
                    else:
                        st.info("📊 Dados insuficientes para evolução. São necessários pelo menos 2 meses.")
                else:
                    st.warning("⚠️ Supabase não conectado.")
            
            st.markdown("---")
            
            # Feedback de Performance
            st.subheader("📝 Feedback de Performance")
            
            col1, col2 = st.columns([2, 1])
            with col1:
                modo_feedback = st.radio(
                    "📋 Formato do Feedback:",
                    options=["SARI - Desenvolvimento", "MIMO - Operacional"],
                    index=0,
                    horizontal=True,
                    key="modo_feedback_individual"
                )
            
            feedback_atual = gerar_feedback_manual(
                analista_selecionado,
                dados,
                media_atendimentos,
                posicao_podio,
                modo_feedback
            )
            
            feedback_editado = st.text_area(
                "✏️ Edite o feedback abaixo:",
                value=feedback_atual,
                height=250,
                key="feedback_editor_individual"
            )
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col1:
                if st.button("🤖 Gerar com IA", use_container_width=True, key="gerar_ia_individual"):
                    with st.spinner("Gerando feedback com IA..."):
                        feedback_gerado = gerar_feedback_ia(
                            analista_selecionado,
                            dados,
                            media_atendimentos,
                            posicao_podio,
                            feedback_editado,
                            modo_feedback
                        )
                        if feedback_gerado:
                            st.session_state[f'feedback_gerado_{analista_selecionado}'] = feedback_gerado
                            st.success("✅ Feedback gerado com IA!")
                            st.rerun()
            
            if st.session_state.get(f'feedback_gerado_{analista_selecionado}'):
                st.markdown("---")
                st.subheader("🤖 Feedback Gerado pela IA")
                st.markdown(st.session_state[f'feedback_gerado_{analista_selecionado}'])
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📋 Usar este feedback", use_container_width=True, key="usar_feedback_individual"):
                        st.session_state[f'feedback_final_{analista_selecionado}'] = st.session_state[f'feedback_gerado_{analista_selecionado}']
                        st.success("✅ Feedback selecionado!")
                        st.rerun()
                with col2:
                    if st.button("🔄 Regenerar", use_container_width=True, key="regenerar_individual"):
                        del st.session_state[f'feedback_gerado_{analista_selecionado}']
                        st.rerun()
            
            feedback_final = st.session_state.get(f'feedback_final_{analista_selecionado}', feedback_editado)
            
            st.markdown("---")
            
            # Análise Técnica em expander
            with st.expander("📊 Análise Técnica de Desempenho", expanded=False):
                analise_tecnica = gerar_analise_tecnica(
                    analista_selecionado,
                    dados,
                    media_atendimentos,
                    podio_efetivo
                )
                st.markdown(analise_tecnica)
            
            st.markdown("---")
            
            # Botões de download
            st.subheader("📄 Download de Relatórios")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**📊 Relatório Individual**")
                st.caption("Apenas a análise técnica (sem feedback)")
                
                relatorio_individual = gerar_relatorio_word(
                    analista_selecionado,
                    dados,
                    analise_tecnica,
                    "",
                    media_atendimentos,
                    podio_efetivo,
                    periodo
                )
                
                st.download_button(
                    label="📥 Baixar Análise (Word)",
                    data=relatorio_individual,
                    file_name=f"Analise_{analista_selecionado.replace(' ', '_')}_{periodo.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                st.markdown("**📄 Relatório Completo**")
                st.caption("Análise técnica + Feedback de Performance (SARI/MIMO)")
                
                relatorio_completo = gerar_relatorio_word(
                    analista_selecionado,
                    dados,
                    analise_tecnica,
                    feedback_final,
                    media_atendimentos,
                    podio_efetivo,
                    periodo
                )
                
                st.download_button(
                    label="📥 Baixar Relatório Completo (Word)",
                    data=relatorio_completo,
                    file_name=f"Relatorio_Completo_{analista_selecionado.replace(' ', '_')}_{periodo.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    
    return resultados

# ============================================
# DASHBOARD COORDENADOR - VERSÃO OTIMIZADA
# ============================================

def dashboard_coordenador_otimizado(periodo, nome_usuario, supabase):
    periodo = padronizar_periodo(periodo)
    
    df_historico = carregar_historico(supabase, mes_ano=periodo)
    
    if df_historico is None or df_historico.empty:
        st.warning(f"⚠️ Nenhum dado encontrado para o período {periodo}")
        return None
    
    resultados = {}
    for _, row in df_historico.iterrows():
        resultados[row['analista']] = {
            'total_atendimentos': row['total_atendimentos'],
            'total_inativos': row['total_inativos'],
            'validos': row['validos'],
            'avaliacoes': row['avaliacoes'],
            'positivos': row['positivos'],
            'negativos': row['negativos'],
            'perc_avaliacoes': row['perc_avaliacoes'],
            'perc_envio': row['perc_envio'],
            'csat': row['csat'],
            'meta_csat': row['meta_csat'],
            'delta_csat': row['delta_csat'],
            'meta_geral': row['meta_geral'],
            'status': row['status'],
            'gestor': row['gestor']
        }
    
    total_analistas = len(resultados)
    total_atendimentos = sum([d['total_atendimentos'] for d in resultados.values()])
    media_atendimentos = total_atendimentos / total_analistas if total_analistas > 0 else 0
    csat_medio = sum([d['csat'] for d in resultados.values()]) / total_analistas if total_analistas > 0 else 0
    perc_avaliacoes_medio = sum([d['perc_avaliacoes'] for d in resultados.values()]) / total_analistas if total_analistas > 0 else 0
    perc_envio_medio = sum([d['perc_envio'] for d in resultados.values()]) / total_analistas if total_analistas > 0 else 0
    metas_superadas = len([d for d in resultados.values() if d['status'] == '🟢 Meta Superada'])
    
    dados_cards = {
        'total_registros': total_analistas,
        'csat_medio': csat_medio,
        'perc_avaliacoes_medio': perc_avaliacoes_medio,
        'perc_envio_medio': perc_envio_medio
    }
    
    st.info(f"📅 {periodo} | 👤 Coordenador: {nome_usuario} | 🏆 Metas: {metas_superadas}/{total_analistas}")
    st.markdown("---")
    
    criar_cards_indicadores_compactos(dados_cards)
    st.markdown("---")
    
    criar_painel_inteligente_compacto(
        csat_medio, perc_avaliacoes_medio, perc_envio_medio,
        metas_superadas, total_analistas
    )
    st.markdown("---")
    
    tab_geral, tab_ranking, tab_analise, tab_individual = st.tabs([
        "📊 Visão Geral",
        "🏆 Rankings",
        "📋 Análise Detalhada",
        "👤 Individual"
    ])
    
    # ===== TAB 1: VISÃO GERAL =====
    with tab_geral:
        st.subheader("📈 Evolução dos Indicadores - Operação")
        df_historico_completo = carregar_historico(supabase)
        
        if df_historico_completo is not None and not df_historico_completo.empty and len(df_historico_completo['mes_ano'].unique()) >= 2:
            df_mensal = df_historico_completo.groupby('mes_ano').agg({
                'csat': 'mean',
                'perc_avaliacoes': 'mean',
                'perc_envio': 'mean'
            }).reset_index()
            
            fig_evolucao = criar_grafico_evolucao_unificado(df_mensal, "Evolução dos Indicadores - Operação")
            if fig_evolucao:
                st.plotly_chart(fig_evolucao, use_container_width=True)
        else:
            st.info("📊 Dados insuficientes para evolução. São necessários pelo menos 2 meses.")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🥇 Top 5 Performers")
            top_performers = calcular_top_performers(resultados, media_atendimentos)
            
            for i, performer in enumerate(top_performers[:5], 1):
                medalha = ["🥇", "🥈", "🥉", "4º", "5º"][i-1]
                cores = ['#FFD700', '#C0C0C0', '#CD7F32', '#8B8B8B', '#A9A9A9']
                st.markdown(f"""
                <div style="background: var(--background-color, #f0f8ff); 
                            padding: 8px; border-radius: 6px; margin-bottom: 4px;
                            border-left: 4px solid {cores[i-1] if i <= 3 else '#8B8B8B'};
                            color: var(--text-color, #333);">
                    <b style="color: {cores[i-1] if i <= 3 else '#8B8B8B'};">{medalha} {performer['analista']}</b>
                    <span style="float: right; color: var(--text-secondary, #666); font-size: 12px;">
                        {performer['status']} | CSAT: {performer['csat']:.1f}%
                    </span>
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            st.subheader("📊 Top 5 Oportunidades")
            oportunidades = calcular_oportunidades(resultados, media_atendimentos)
            
            for i, oport in enumerate(oportunidades[:5], 1):
                cor = "#e74c3c" if oport['csat'] < 85 else "#f39c12"
                st.markdown(f"""
                <div style="background: var(--background-color, #fff5f5); 
                            padding: 8px; border-radius: 6px; margin-bottom: 4px;
                            border-left: 4px solid {cor};
                            color: var(--text-color, #333);">
                    <b>{i}º {oport['analista']}</b>
                    <span style="float: right; color: var(--text-secondary, #666); font-size: 12px;">
                        CSAT: {oport['csat']:.1f}%
                    </span>
                </div>
                """, unsafe_allow_html=True)
    
    # ===== TAB 2: RANKINGS =====
    with tab_ranking:
        st.subheader("🏆 Pódio do Mês - Operação")
        podio = calcular_podio(resultados, media_atendimentos)
        mostrar_podio(podio)
        
        st.markdown("---")
        st.subheader("📊 Ranking Completo - Operação")
        dados_ranking = []
        for i, (analista, dados) in enumerate(sorted(resultados.items(), key=lambda x: x[1]['csat'], reverse=True), 1):
            dados_ranking.append({
                'Posição': i,
                'Analista': analista,
                'Gestor': dados['gestor'],
                'CSAT': f"{dados['csat']:.1f}%",
                'Avaliações': f"{dados['perc_avaliacoes']:.1f}%",
                'Atendimentos': dados['total_atendimentos'],
                'Status': dados['status']
            })
        st.dataframe(pd.DataFrame(dados_ranking), use_container_width=True, hide_index=True)
    
    # ===== TAB 3: ANÁLISE DETALHADA =====
    with tab_analise:
        st.subheader("📋 Análise de Desempenho - Operação")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            filtro_gestor = st.selectbox("Gestor", ["Todos"] + GESTORES_VALIDOS)
        with col2:
            filtro_status = st.selectbox("Status", ["Todos", "🟢 Meta Superada", "🟡 Atenção", "🔴 Crítico"])
        with col3:
            filtro_analista = st.selectbox("Analista", ["Todos"] + list(resultados.keys()))
        with col4:
            filtro_ordenacao = st.selectbox("Ordenar por", ["CSAT (decrescente)", "CSAT (crescente)", "Avaliações", "Atendimentos"])
        
        dados_filtrados = resultados
        if filtro_gestor != "Todos":
            dados_filtrados = {k: v for k, v in dados_filtrados.items() if v['gestor'] == filtro_gestor}
        if filtro_status != "Todos":
            dados_filtrados = {k: v for k, v in dados_filtrados.items() if v['status'] == filtro_status}
        if filtro_analista != "Todos":
            dados_filtrados = {k: v for k, v in dados_filtrados.items() if k == filtro_analista}
        
        if filtro_ordenacao == "CSAT (decrescente)":
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['csat'], reverse=True)
        elif filtro_ordenacao == "CSAT (crescente)":
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['csat'])
        elif filtro_ordenacao == "Avaliações":
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['perc_avaliacoes'], reverse=True)
        else:
            dados_ordenados = sorted(dados_filtrados.items(), key=lambda x: x[1]['total_atendimentos'], reverse=True)
        
        dados_tabela = []
        for analista, dados in dados_ordenados:
            delta_avaliacoes = dados['perc_avaliacoes'] - dados['meta_geral']
            dados_tabela.append({
                'Analista': analista,
                'Gestor': dados['gestor'],
                'CSAT': f"{dados['csat']:.1f}%",
                'Meta CSAT': f"{dados['meta_csat']}%",
                'Delta CSAT': f"{dados['delta_csat']:+.1f}%",
                '% Avaliações': f"{dados['perc_avaliacoes']:.1f}%",
                'Meta Avaliações': f"{dados['meta_geral']}%",
                'Delta Avaliações': f"{delta_avaliacoes:+.1f}%",
                '% Envio': f"{dados['perc_envio']:.1f}%",
                '💬 Atendimentos': dados['total_atendimentos'],
                'Status': dados['status']
            })
        
        if dados_tabela:
            st.dataframe(pd.DataFrame(dados_tabela), use_container_width=True, hide_index=True)
            
            st.markdown("---")
            st.subheader("📊 Resumo Estatístico - Operação")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Média Atendimentos", f"{media_atendimentos:.0f}")
            with col2:
                st.metric("Maior Atendimentos", f"{max([d['total_atendimentos'] for d in resultados.values()])}")
            with col3:
                st.metric("Menor Atendimentos", f"{min([d['total_atendimentos'] for d in resultados.values()])}")
            with col4:
                st.metric("Total Atendimentos", f"{total_atendimentos}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Média CSAT", f"{csat_medio:.1f}%")
            with col2:
                st.metric("Média Avaliações", f"{perc_avaliacoes_medio:.1f}%")
            with col3:
                st.metric("Média Envio", f"{perc_envio_medio:.1f}%")
        else:
            st.info("Nenhum dado encontrado com os filtros selecionados.")
    
    # ===== TAB 4: INDIVIDUAL =====
    with tab_individual:
        st.subheader("👤 Relatório Individual - Operação")
        
        analista_selecionado = st.selectbox(
            "Selecione o Analista",
            list(resultados.keys()),
            key="analista_individual_coord"
        )
        
        if analista_selecionado:
            dados = resultados[analista_selecionado]
            podio = calcular_podio(resultados, media_atendimentos)
            
            posicao_podio = None
            for i, (nome, _, _, _) in enumerate(podio, 1):
                if nome == analista_selecionado:
                    posicao_podio = i
                    break
            
            st.markdown(f"### 📊 {analista_selecionado}")
            st.markdown(f"**Gestor:** {dados['gestor']}")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                cor_csat = "#2ecc71" if dados['csat'] >= dados['meta_csat'] else "#e74c3c"
                criar_card_compacto(
                    "⭐ CSAT",
                    f"{dados['csat']:.1f}%",
                    f"Meta: {dados['meta_csat']}%",
                    f"{dados['delta_csat']:+.1f}%",
                    cor_csat
                )
            
            with col2:
                diff_avaliacoes = dados['perc_avaliacoes'] - dados['meta_geral']
                cor_avaliacoes = "#2ecc71" if diff_avaliacoes >= 0 else "#e74c3c"
                criar_card_compacto(
                    "📝 Avaliações",
                    f"{dados['perc_avaliacoes']:.1f}%",
                    f"Meta: {dados['meta_geral']}%",
                    f"{diff_avaliacoes:+.1f}%",
                    cor_avaliacoes
                )
            
            with col3:
                diff_envio = dados['perc_envio'] - 80
                cor_envio = "#2ecc71" if diff_envio >= 0 else "#f39c12"
                criar_card_compacto(
                    "📤 Envio",
                    f"{dados['perc_envio']:.1f}%",
                    "Meta: 80%",
                    f"{diff_envio:+.1f}%",
                    cor_envio
                )
            
            with col4:
                diff_atend = dados['total_atendimentos'] - media_atendimentos
                cor_atend = "#2ecc71" if diff_atend >= 0 else "#e74c3c"
                criar_card_compacto(
                    "💬 Atendimentos",
                    f"{dados['total_atendimentos']}",
                    f"Média: {media_atendimentos:.0f}",
                    f"{diff_atend:+.0f}",
                    cor_atend
                )
            
            st.markdown(f"**Status:** {dados['status']}")
            if posicao_podio:
                medalhas = ["🥇", "🥈", "🥉"]
                st.markdown(f"**Posição:** {medalhas[posicao_podio-1]} {posicao_podio}º lugar no pódio")
            
            st.markdown("---")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📊 Radar de Performance")
                categorias = ['CSAT', '% Avaliações', '% Envio']
                valores = [dados['csat'], dados['perc_avaliacoes'], dados['perc_envio']]
                metas = [dados['meta_csat'], dados['meta_geral'], 80]
                
                fig_radar = go.Figure()
                fig_radar.add_trace(go.Scatterpolar(
                    r=valores,
                    theta=categorias,
                    fill='toself',
                    name=analista_selecionado,
                    line_color='#2ecc71',
                    fillcolor='rgba(46, 204, 113, 0.3)'
                ))
                fig_radar.add_trace(go.Scatterpolar(
                    r=metas,
                    theta=categorias,
                    fill='toself',
                    name='Meta',
                    line_color='#e74c3c',
                    fillcolor='rgba(231, 76, 60, 0.1)',
                    line_dash='dash'
                ))
                fig_radar.update_layout(
                    height=380,
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 100],
                            tickfont=dict(color='var(--text-color, #e0e0e0)')
                        ),
                        angularaxis=dict(
                            tickfont=dict(color='var(--text-color, #e0e0e0)')
                        )
                    ),
                    showlegend=True,
                    legend=dict(
                        orientation='h',
                        yanchor='bottom',
                        y=-0.15,
                        xanchor='center',
                        x=0.5
                    ),
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='var(--text-color, #e0e0e0)')
                )
                st.plotly_chart(fig_radar, use_container_width=True)
            
            with col2:
                st.subheader("📈 Evolução Mensal")
                
                if supabase:
                    df_historico_analista = carregar_historico(supabase, analista=analista_selecionado)
                    if df_historico_analista is not None and not df_historico_analista.empty and len(df_historico_analista['mes_ano'].unique()) >= 2:
                        fig_individual = gerar_grafico_mensal(
                            analista_selecionado,
                            df_historico_analista,
                            dados['meta_csat'],
                            dados['meta_geral']
                        )
                        if fig_individual:
                            st.plotly_chart(fig_individual, use_container_width=True)
                        else:
                            st.info("📊 Dados insuficientes para gráfico individual.")
                    else:
                        st.info("📊 Dados insuficientes para evolução. São necessários pelo menos 2 meses.")
                else:
                    st.warning("⚠️ Supabase não conectado.")
            
            st.markdown("---")
            
            # Feedback
            st.subheader("📝 Feedback de Performance")
            
            col1, col2 = st.columns([2, 1])
            with col1:
                modo_feedback = st.radio(
                    "📋 Formato do Feedback:",
                    options=["SARI - Desenvolvimento", "MIMO - Operacional"],
                    index=0,
                    horizontal=True,
                    key="modo_feedback_coord"
                )
            
            feedback_atual = gerar_feedback_manual(
                analista_selecionado,
                dados,
                media_atendimentos,
                posicao_podio,
                modo_feedback
            )
            
            feedback_editado = st.text_area(
                "✏️ Edite o feedback abaixo:",
                value=feedback_atual,
                height=250,
                key="feedback_editor_coord"
            )
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col1:
                if st.button("🤖 Gerar com IA", use_container_width=True, key="gerar_ia_coord"):
                    with st.spinner("Gerando feedback com IA..."):
                        feedback_gerado = gerar_feedback_ia(
                            analista_selecionado,
                            dados,
                            media_atendimentos,
                            posicao_podio,
                            feedback_editado,
                            modo_feedback
                        )
                        if feedback_gerado:
                            st.session_state[f'feedback_gerado_coord_{analista_selecionado}'] = feedback_gerado
                            st.success("✅ Feedback gerado com IA!")
                            st.rerun()
            
            if st.session_state.get(f'feedback_gerado_coord_{analista_selecionado}'):
                st.markdown("---")
                st.subheader("🤖 Feedback Gerado pela IA")
                st.markdown(st.session_state[f'feedback_gerado_coord_{analista_selecionado}'])
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📋 Usar este feedback", use_container_width=True, key="usar_feedback_coord"):
                        st.session_state[f'feedback_final_coord_{analista_selecionado}'] = st.session_state[f'feedback_gerado_coord_{analista_selecionado}']
                        st.success("✅ Feedback selecionado!")
                        st.rerun()
                with col2:
                    if st.button("🔄 Regenerar", use_container_width=True, key="regenerar_coord"):
                        del st.session_state[f'feedback_gerado_coord_{analista_selecionado}']
                        st.rerun()
            
            feedback_final = st.session_state.get(f'feedback_final_coord_{analista_selecionado}', feedback_editado)
            
            st.markdown("---")
            
            with st.expander("📊 Análise Técnica de Desempenho", expanded=False):
                analise_tecnica = gerar_analise_tecnica(
                    analista_selecionado,
                    dados,
                    media_atendimentos,
                    podio
                )
                st.markdown(analise_tecnica)
            
            st.markdown("---")
            
            st.subheader("📄 Download de Relatórios")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**📊 Relatório Individual**")
                st.caption("Apenas a análise técnica (sem feedback)")
                
                relatorio_individual = gerar_relatorio_word(
                    analista_selecionado,
                    dados,
                    analise_tecnica,
                    "",
                    media_atendimentos,
                    podio,
                    periodo
                )
                
                st.download_button(
                    label="📥 Baixar Análise (Word)",
                    data=relatorio_individual,
                    file_name=f"Analise_{analista_selecionado.replace(' ', '_')}_{periodo.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                st.markdown("**📄 Relatório Completo**")
                st.caption("Análise técnica + Feedback de Performance (SARI/MIMO)")
                
                relatorio_completo = gerar_relatorio_word(
                    analista_selecionado,
                    dados,
                    analise_tecnica,
                    feedback_final,
                    media_atendimentos,
                    podio,
                    periodo
                )
                
                st.download_button(
                    label="📥 Baixar Relatório Completo (Word)",
                    data=relatorio_completo,
                    file_name=f"Relatorio_Completo_{analista_selecionado.replace(' ', '_')}_{periodo.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    
    return resultados

# ============================================
# PAGINAS DO SISTEMA
# ============================================

def pagina_importar_periodo():
    st.header("📁 Importar Período")
    
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não configurado")
        return
    
    analistas_config = carregar_analistas()
    gestor_ativo = st.session_state.get('gestor', None)
    
    col1, col2 = st.columns(2)
    with col1:
        arquivo_satisfacao = st.file_uploader("Arquivo de Satisfação", type=['xlsx', 'csv'], key="import_satisfacao")
    with col2:
        arquivo_inativos = st.file_uploader("Arquivo de Inatividade", type=['xlsx', 'csv'], key="import_inativos")
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    with col1:
        meses = ['Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho', 
                 'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro']
        mes_selecionado = st.selectbox("Mês", meses, key="import_mes")
    with col2:
        anos = list(range(2020, 2031))
        ano_selecionado = st.selectbox("Ano", anos, index=anos.index(2026), key="import_ano")
    
    periodo = padronizar_periodo(f"{mes_selecionado} {ano_selecionado}")
    st.info(f"📅 Período selecionado: **{periodo}**")
    
    if st.button("🚀 Processar Dados", use_container_width=True, key="import_processar"):
        if not arquivo_satisfacao or not arquivo_inativos:
            st.error("❌ Envie os dois arquivos.")
        else:
            with st.spinner("Processando dados..."):
                try:
                    df_satisfacao = carregar_arquivo_satisfacao(arquivo_satisfacao)
                    if df_satisfacao is None:
                        st.stop()
                    df_inativos = carregar_arquivo_inativos(arquivo_inativos)
                    if df_inativos is None:
                        st.stop()
                    
                    gestor_salvar = st.session_state.get('gestor', GESTOR_MARCOS)
                    
                    resultados_temp = processar_dados(df_satisfacao, df_inativos, analistas_config)
                    dados_por_gestor = {}
                    for analista, d in resultados_temp.items():
                        gestor_real = d.get('gestor', gestor_salvar)
                        if gestor_real not in dados_por_gestor:
                            dados_por_gestor[gestor_real] = []
                        dados_por_gestor[gestor_real].append((analista, d))
                    
                    existe_duplicidade = False
                    gestores_com_dados = []
                    for gestor_real, _ in dados_por_gestor.items():
                        if verificar_periodo_existente(supabase, periodo, gestor_real):
                            existe_duplicidade = True
                            gestores_com_dados.append(gestor_real)
                    
                    if existe_duplicidade:
                        st.warning(f"⚠️ Já existe relatório para {periodo} do(s) gestor(es): {', '.join(gestores_com_dados)}")
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("❌ Cancelar", use_container_width=True):
                                st.stop()
                        with col2:
                            if st.button("🔄 Substituir", use_container_width=True):
                                resultados = processar_dados(df_satisfacao, df_inativos, analistas_config)
                                st.session_state.resultados = resultados
                                st.session_state.processado = True
                                st.session_state.periodo = periodo
                                sucesso, mensagem = substituir_historico(supabase, resultados, periodo, gestor_salvar)
                                if sucesso:
                                    st.success(f"✅ Período {periodo} substituído!")
                                    st.rerun()
                                else:
                                    st.error(f"❌ Erro: {mensagem}")
                        st.stop()
                    
                    resultados = processar_dados(df_satisfacao, df_inativos, analistas_config)
                    st.session_state.resultados = resultados
                    st.session_state.processado = True
                    st.session_state.periodo = periodo
                    
                    sucesso, mensagem = salvar_historico(supabase, resultados, periodo, gestor_salvar)
                    if sucesso:
                        st.success(f"✅ Dados salvos! Período: {periodo}")
                        st.rerun()
                    else:
                        st.warning(f"⚠️ Dados NÃO salvos: {mensagem}")
                except Exception as e:
                    st.error(f"❌ Erro: {str(e)}")

def pagina_gerenciar_periodos():
    st.header("📂 Gerenciar Períodos")
    
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não configurado")
        return
    
    acesso_total = st.session_state.get('acesso_total', False)
    gestor_ativo = st.session_state.get('gestor', None)
    
    if acesso_total:
        periodos = listar_periodos(supabase)
    else:
        periodos = listar_periodos(supabase, gestor_ativo)
    
    if periodos:
        dados_tabela = []
        for p in periodos:
            df_periodo = carregar_historico(supabase, mes_ano=p['mes_ano'], gestor=p['gestor'])
            qtd_analistas = len(df_periodo) if df_periodo is not None else 0
            dados_tabela.append({
                'Período': p['mes_ano'],
                'Gestor': p['gestor'],
                'Analistas': qtd_analistas,
                'Data Importação': p.get('data_criacao', 'N/A')[:10] if p.get('data_criacao') else 'N/A'
            })
        st.dataframe(pd.DataFrame(dados_tabela), use_container_width=True, hide_index=True)
        st.markdown("---")
        
        st.subheader("🗑️ Excluir Período")
        col1, col2 = st.columns([2, 1])
        with col1:
            periodo_para_excluir = st.selectbox("Selecione o período", [p['mes_ano'] for p in periodos], key="periodo_excluir")
        with col2:
            confirmar = st.checkbox("✅ Confirmar exclusão", key="confirmar_excluir")
        
        if st.button("🗑️ Excluir Período", use_container_width=True, key="btn_excluir_periodo"):
            if not confirmar:
                st.error("⚠️ Marque a caixa de confirmação.")
            else:
                gestor_periodo = next((p['gestor'] for p in periodos if p['mes_ano'] == periodo_para_excluir), None)
                if gestor_periodo:
                    if not acesso_total and gestor_periodo != gestor_ativo:
                        st.error("❌ Sem permissão para excluir.")
                    else:
                        sucesso, mensagem = excluir_periodo(supabase, periodo_para_excluir, gestor_periodo)
                        if sucesso:
                            st.success(f"✅ {mensagem}")
                            st.rerun()
                        else:
                            st.error(f"❌ Erro: {mensagem}")
    else:
        if acesso_total:
            st.info("Nenhum período importado na operação.")
        else:
            st.info(f"Nenhum período importado para {gestor_ativo}.")

def pagina_historico():
    st.header("📈 Histórico")
    
    supabase = init_supabase()
    if not supabase:
        st.error("❌ Supabase não configurado")
        return
    
    acesso_total = st.session_state.get('acesso_total', False)
    gestor_ativo = st.session_state.get('gestor', None)
    
    if acesso_total:
        df_historico = carregar_historico(supabase)
    else:
        df_historico = carregar_historico(supabase, gestor=gestor_ativo)
    
    if df_historico is not None and not df_historico.empty:
        col1, col2, col3 = st.columns(3)
        with col1:
            meses_disponiveis = ['Todos'] + sorted(df_historico['mes_ano'].unique().tolist())
            mes_selecionado = st.selectbox("Selecione o Mês", meses_disponiveis, key="hist_mes")
        with col2:
            if acesso_total:
                gestores_disponiveis = ['Todos'] + sorted(df_historico['gestor'].unique().tolist())
            else:
                gestores_disponiveis = [gestor_ativo]
            gestor_filtro = st.selectbox("Selecione o Gestor", gestores_disponiveis, key="hist_gestor")
        with col3:
            if acesso_total and gestor_filtro != 'Todos':
                df_analistas = df_historico[df_historico['gestor'] == gestor_filtro]
            elif not acesso_total:
                df_analistas = df_historico[df_historico['gestor'] == gestor_ativo]
            else:
                df_analistas = df_historico
            analistas_disponiveis = ['Todos'] + sorted(df_analistas['analista'].unique().tolist())
            analista_filtro = st.selectbox("Selecione o Analista", analistas_disponiveis, key="hist_analista")
        
        df_filtrado = df_historico.copy()
        if mes_selecionado != 'Todos':
            df_filtrado = df_filtrado[df_filtrado['mes_ano'] == mes_selecionado]
        if gestor_filtro != 'Todos' and acesso_total:
            df_filtrado = df_filtrado[df_filtrado['gestor'] == gestor_filtro]
        elif not acesso_total:
            df_filtrado = df_filtrado[df_filtrado['gestor'] == gestor_ativo]
        if analista_filtro != 'Todos':
            df_filtrado = df_filtrado[df_filtrado['analista'] == analista_filtro]
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Registros", len(df_filtrado))
        with col2:
            st.metric("CSAT Médio", f"{df_filtrado['csat'].mean():.2f}%")
        with col3:
            st.metric("% Avaliações Médio", f"{df_filtrado['perc_avaliacoes'].mean():.2f}%")
        with col4:
            st.metric("% Envio Médio", f"{df_filtrado['perc_envio'].mean():.2f}%")
        
        st.markdown("---")
        
        if not df_filtrado.empty:
            df_agrupado = df_filtrado.groupby('mes_ano').agg({
                'csat': 'mean',
                'perc_avaliacoes': 'mean',
                'perc_envio': 'mean'
            }).reset_index()
            
            fig_evolucao = criar_grafico_evolucao_unificado(df_agrupado, "📈 Evolução Mensal - Histórico")
            if fig_evolucao:
                st.plotly_chart(fig_evolucao, use_container_width=True)
            
            if analista_filtro != 'Todos':
                st.subheader(f"📊 Evolução Mensal - {analista_filtro}")
                df_analista = df_filtrado[df_filtrado['analista'] == analista_filtro]
                if not df_analista.empty:
                    meta_csat = df_analista['meta_csat'].iloc[0] if 'meta_csat' in df_analista.columns else 90
                    meta_geral = df_analista['meta_geral'].iloc[0] if 'meta_geral' in df_analista.columns else 25
                    fig_mensal = gerar_grafico_mensal(analista_filtro, df_historico, meta_csat, meta_geral)
                    if fig_mensal:
                        st.plotly_chart(fig_mensal, use_container_width=True)
        
        st.subheader("📋 Dados Históricos")
        st.dataframe(df_filtrado, use_container_width=True)
    else:
        if acesso_total:
            st.warning("Nenhum dado histórico encontrado na operação.")
        else:
            st.warning(f"Nenhum dado histórico encontrado para {gestor_ativo}.")

def pagina_gerenciar_usuarios():
    st.header("👥 Gerenciar Usuários")
    gerenciar_usuarios_supabase()

def pagina_gerenciar_analistas():
    st.header("📝 Gerenciar Analistas")
    
    analistas_config = carregar_analistas()
    
    gestor_selecionado = st.selectbox("Selecione o Gestor", list(analistas_config.keys()), key="gerenciar_gestor")
    if not gestor_selecionado:
        return
    
    config = analistas_config[gestor_selecionado]
    
    nova_meta = st.number_input(
        "Meta Geral de Avaliações (%)", 
        min_value=0, max_value=100, 
        value=config['meta_geral_avaliacoes'],
        key="meta_geral_input"
    )
    if nova_meta != config['meta_geral_avaliacoes']:
        config['meta_geral_avaliacoes'] = nova_meta
        salvar_analistas(analistas_config)
    
    st.markdown("---")
    
    st.subheader("🔄 Mover Analista entre Times")
    todos_analistas = []
    for gestor, cfg in analistas_config.items():
        for analista, dados in cfg['membros'].items():
            if dados['ativo']:
                todos_analistas.append((analista, gestor))
    
    if todos_analistas:
        col1, col2, col3 = st.columns([2, 2, 1])
        with col1:
            analista_para_mover = st.selectbox("Selecione o analista", [a[0] for a in todos_analistas], key="mover_analista")
        with col2:
            gestor_atual = next((g for a, g in todos_analistas if a == analista_para_mover), None)
            todos_gestores = list(analistas_config.keys())
            indice_atual = todos_gestores.index(gestor_atual) if gestor_atual in todos_gestores else 0
            novo_gestor = st.selectbox(f"Time atual: {gestor_atual}", todos_gestores, index=indice_atual, key="novo_gestor")
        with col3:
            if novo_gestor and novo_gestor != gestor_atual:
                if st.button("🔄 Mover", use_container_width=True, key="btn_mover_analista"):
                    meta_atual = analistas_config[gestor_atual]['membros'][analista_para_mover]['meta_csat']
                    del analistas_config[gestor_atual]['membros'][analista_para_mover]
                    analistas_config[novo_gestor]['membros'][analista_para_mover] = {"meta_csat": meta_atual, "ativo": True}
                    salvar_analistas(analistas_config)
                    st.success(f"✅ {analista_para_mover} movido para {novo_gestor}!")
                    st.rerun()
            elif novo_gestor == gestor_atual:
                st.info("ℹ️ O analista já está neste time.")
    else:
        st.info("Nenhum analista ativo para mover.")
    
    st.markdown("---")
    
    st.subheader(f"👥 Membros do Time: {gestor_selecionado}")
    membros = config['membros']
    
    col1, col2, col3 = st.columns([3, 1, 1])
    with col1:
        novo_nome = st.text_input("Nome do novo analista", key="novo_analista_nome")
    with col2:
        nova_meta_csat = st.number_input("Meta CSAT", min_value=0, max_value=100, value=86, key="novo_analista_meta")
    with col3:
        if st.button("➕ Adicionar", use_container_width=True, key="btn_adicionar_analista") and novo_nome:
            if novo_nome not in membros:
                membros[novo_nome] = {"meta_csat": nova_meta_csat, "ativo": True}
                salvar_analistas(analistas_config)
                st.rerun()
            else:
                st.error("❌ Analista já existe!")
    
    st.markdown("---")
    
    dados_tabela = []
    for nome, dados in membros.items():
        dados_tabela.append({
            "Analista": nome, 
            "Meta CSAT": f"{dados['meta_csat']}%", 
            "Status": "✅ Ativo" if dados['ativo'] else "❌ Inativo"
        })
    if dados_tabela:
        df_membros = pd.DataFrame(dados_tabela)
        st.dataframe(df_membros, use_container_width=True, hide_index=True)
    
    st.subheader("✏️ Editar/Remover Membros")
    membro_selecionado = st.selectbox("Selecione um membro para editar", list(membros.keys()), key="editar_membro")
    if membro_selecionado:
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            nova_meta_membro = st.number_input(
                "Nova Meta CSAT", 
                min_value=0, max_value=100, 
                value=membros[membro_selecionado]['meta_csat'],
                key="editar_meta"
            )
        with col2:
            novo_status = st.checkbox("Ativo", value=membros[membro_selecionado]['ativo'], key="editar_status")
        with col3:
            if st.button("🗑️ Remover", use_container_width=True, key="btn_remover_analista"):
                del membros[membro_selecionado]
                salvar_analistas(analistas_config)
                st.rerun()
        
        if nova_meta_membro != membros[membro_selecionado]['meta_csat'] or novo_status != membros[membro_selecionado]['ativo']:
            membros[membro_selecionado]['meta_csat'] = nova_meta_membro
            membros[membro_selecionado]['ativo'] = novo_status
            salvar_analistas(analistas_config)
            st.success("✅ Alterações salvas!")
            st.rerun()

def pagina_configuracoes():
    st.header("⚙️ Configurações do Sistema")
    
    supabase = init_supabase()
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Resetar Carine", use_container_width=True):
            resetar_usuario_carine()
            st.rerun()
        if st.button("🔄 Resetar Marcos", use_container_width=True):
            resetar_usuario_marcos()
            st.rerun()
        if st.button("🔄 Resetar Polyana", use_container_width=True):
            resetar_usuario_polyana()
            st.rerun()
        if st.button("📥 Adicionar TODOS os Analistas da Polyana", use_container_width=True):
            adicionar_dados_teste_polyana()
            st.rerun()
    with col2:
        if st.button("🧹 LIMPAR CACHE E FORÇAR PERFIL", use_container_width=True):
            limpar_cache_completo()
            st.rerun()
        if st.button("🔍 Diagnóstico do Sistema", use_container_width=True):
            diagnosticar_sistema()
    
    st.markdown("---")
    st.subheader("💾 Backup e Restore")
    
    if supabase:
        buffer, mensagem = gerar_backup_excel(supabase)
        if buffer:
            data_atual = datetime.now().strftime("%Y%m%d_%H%M%S")
            nome_arquivo = f"backup_sistema_{data_atual}.xlsx"
            
            st.download_button(
                label="📥 Baixar Backup (Excel)",
                data=buffer,
                file_name=nome_arquivo,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        else:
            st.warning(mensagem)
    else:
        st.error("❌ Supabase não configurado.")
    
    st.markdown("---")
    
    st.subheader("📤 Restaurar Backup")
    st.caption("⚠️ Restaure apenas com dados do mesmo sistema para evitar conflitos.")
    
    arquivo_backup = st.file_uploader(
        "Selecione o arquivo de backup (.xlsx)",
        type=['xlsx'],
        key="upload_backup_config"
    )
    
    modo_restore = st.radio(
        "Modo de restauração:",
        ["Mesclar (recomendado, não apaga nada)", "Substituir tudo (APAGA os dados atuais antes de restaurar)"],
        key="modo_restore_config"
    )
    
    confirmacao_ok = True
    if "Substituir" in modo_restore:
        confirmacao = st.text_input(
            "Digite CONFIRMAR para prosseguir com a substituição:",
            type="password",
            key="confirmacao_substituir_config"
        )
        if confirmacao != "CONFIRMAR":
            confirmacao_ok = False
            st.warning("⚠️ Digite CONFIRMAR exatamente como está escrito para habilitar a restauração.")
    
    if arquivo_backup and confirmacao_ok and supabase:
        if st.button("🔄 Restaurar Dados", use_container_width=True, key="btn_restaurar_config"):
            with st.spinner("Restaurando dados..."):
                modo = "substituir" if "Substituir" in modo_restore else "mesclar"
                sucesso, mensagem = restaurar_backup_excel(supabase, arquivo_backup, modo)
                if sucesso:
                    st.success(f"✅ {mensagem}")
                    st.rerun()
                else:
                    st.error(f"❌ {mensagem}")
    elif not supabase:
        st.error("❌ Supabase não configurado.")

# ============================================
# DASHBOARD PRINCIPAL
# ============================================

def dashboard_principal():
    supabase = init_supabase()
    acesso_total = st.session_state.get('acesso_total', False)
    gestor_ativo = st.session_state.get('gestor', None)
    nome_usuario = st.session_state.get('nome_usuario', '')
    
    if supabase:
        if acesso_total:
            periodos_disponiveis = listar_periodos(supabase)
        else:
            periodos_disponiveis = listar_periodos(supabase, gestor_ativo)
        
        if periodos_disponiveis:
            periodos_lista = ordenar_meses([p['mes_ano'] for p in periodos_disponiveis])[::-1]
            
            periodo_atual = st.session_state.get('periodo', periodos_lista[0] if periodos_lista else None)
            
            if periodo_atual not in periodos_lista and periodos_lista:
                periodo_atual = periodos_lista[0]
            
            periodo_selecionado = st.selectbox(
                "📅 Selecione o Período",
                periodos_lista,
                index=periodos_lista.index(periodo_atual) if periodo_atual in periodos_lista else 0,
                key="seletor_periodo_dashboard"
            )
            
            if periodo_selecionado != periodo_atual:
                st.session_state.periodo = periodo_selecionado
                st.rerun()
    
    if not st.session_state.get('processado', False):
        if supabase and periodos_disponiveis:
            ultimo_periodo = ordenar_meses([p['mes_ano'] for p in periodos_disponiveis])[::-1][0]
            
            df_historico = carregar_historico(supabase, mes_ano=ultimo_periodo)
            if df_historico is not None and not df_historico.empty:
                resultados = {}
                for _, row in df_historico.iterrows():
                    resultados[row['analista']] = {
                        'total_atendimentos': row['total_atendimentos'],
                        'total_inativos': row['total_inativos'],
                        'validos': row['validos'],
                        'avaliacoes': row['avaliacoes'],
                        'positivos': row['positivos'],
                        'negativos': row['negativos'],
                        'perc_avaliacoes': row['perc_avaliacoes'],
                        'perc_envio': row['perc_envio'],
                        'csat': row['csat'],
                        'meta_csat': row['meta_csat'],
                        'delta_csat': row['delta_csat'],
                        'meta_geral': row['meta_geral'],
                        'status': row['status'],
                        'gestor': row['gestor']
                    }
                st.session_state.resultados = resultados
                st.session_state.processado = True
                st.session_state.periodo = ultimo_periodo
    
    if not st.session_state.get('processado', False) or not st.session_state.get('resultados'):
        st.info("📊 Nenhum dado carregado. Faça a importação de um período.")
        return
    
    periodo = st.session_state.get('periodo', datetime.now().strftime('%B %Y'))
    dashboard_resultados = st.session_state.resultados
    
    if acesso_total:
        st.info("🔑 Perfil: Coordenador - Visualizando toda a operação")
        
        if 'visao_coordenador' not in st.session_state:
            st.session_state.visao_coordenador = "🏢 Visão Geral (Toda a Operação)"
        
        opcoes_visao = ["🏢 Visão Geral (Toda a Operação)"]
        for gestor in GESTORES_VALIDOS:
            opcoes_visao.append(f"👥 {gestor}")
        
        visao_selecionada = st.radio(
            "📊 Visualizar:",
            opcoes_visao,
            index=opcoes_visao.index(st.session_state.visao_coordenador) if st.session_state.visao_coordenador in opcoes_visao else 0,
            horizontal=True,
            key="seletor_visao_coordenador"
        )
        
        st.session_state.visao_coordenador = visao_selecionada
        
        if visao_selecionada == "🏢 Visão Geral (Toda a Operação)":
            dashboard_resultados = dashboard_coordenador_otimizado(periodo, nome_usuario, supabase)
        else:
            gestor_escolhido = visao_selecionada.replace("👥 ", "")
            st.info(f"👥 Visualizando time: {gestor_escolhido}")
            dashboard_resultados = dashboard_gestor_otimizado(periodo, gestor_escolhido, supabase)
    else:
        st.info(f"👥 Perfil: Gestor - Visualizando equipe: {gestor_ativo}")
        dashboard_resultados = dashboard_gestor_otimizado(periodo, gestor_ativo, supabase)
    
    return dashboard_resultados

# ============================================
# MAIN
# ============================================

def main():
    st.title("📊 Sistema de Performance - Relatórios Automáticos")
    st.markdown("---")
    
    if not fazer_login_corrigido():
        st.info("👋 Faça login na barra lateral para acessar o sistema.")
        return
    
    forcar_perfil_correto()
    
    st.sidebar.markdown("---")
    st.sidebar.header("📋 Navegação")
    
    menu_opcoes = [
        "📊 Dashboard",
        "📁 Importar Período",
        "📂 Gerenciar Períodos",
        "📈 Histórico",
        "👥 Gerenciar Usuários",
        "📝 Gerenciar Analistas",
        "⚙️ Configurações"
    ]
    
    if 'pagina_atual' not in st.session_state:
        st.session_state.pagina_atual = "📊 Dashboard"
    
    for opcao in menu_opcoes:
        if st.sidebar.button(opcao, use_container_width=True, key=f"menu_{opcao}"):
            st.session_state.pagina_atual = opcao
            st.rerun()
    
    pagina = st.session_state.pagina_atual
    
    if pagina == "📊 Dashboard":
        dashboard_principal()
    elif pagina == "📁 Importar Período":
        pagina_importar_periodo()
    elif pagina == "📂 Gerenciar Períodos":
        pagina_gerenciar_periodos()
    elif pagina == "📈 Histórico":
        pagina_historico()
    elif pagina == "👥 Gerenciar Usuários":
        pagina_gerenciar_usuarios()
    elif pagina == "📝 Gerenciar Analistas":
        pagina_gerenciar_analistas()
    elif pagina == "⚙️ Configurações":
        pagina_configuracoes()
    
    st.markdown("---")
    st.markdown('<div style="text-align: center; color: var(--text-secondary, #666); font-size: 12px;">Sistema de Performance v15.0 | Versão Completa Otimizada</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
